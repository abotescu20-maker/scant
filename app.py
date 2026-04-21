"""
Insurance Broker AI Assistant — Chainlit UI
============================================
Chat interface for non-technical insurance broker employees.
Uses proprietary AI engine.
Features: PDF/image upload (AI analysis), email offers, export PDF/XLSX/DOCX
"""
import sys
import os
import io
import json
import sqlite3
import tempfile
from pathlib import Path
from datetime import date, timedelta
from dotenv import load_dotenv

import chainlit as cl
import anthropic
import base64

# ── Admin DB helpers (imported only if tables exist) ──────────────────────────
try:
    from shared.db import (
        get_user_by_email, get_user_tools, log_audit, record_token_usage,
        init_admin_tables,
        # Projects & persistent conversations
        create_project, list_projects,
        create_conversation, list_conversations,
        update_conversation_title,
        save_conversation_history, load_conversation_history,
        # Client-linked conversations
        set_conversation_client,
        list_clients_with_conversations,
        list_conversations_for_client,
        get_all_clients_for_picker,
        # Conversation management
        delete_conversation, search_conversations,
    )
    from shared.auth import verify_password
    init_admin_tables()
    ADMIN_ENABLED = True
except Exception:
    ADMIN_ENABLED = False

# ── Session key constants ──────────────────────────────────────────────────────
_SK_HISTORY    = "history"
_SK_CONV_ID    = "conversation_id"
_SK_PROJECT_ID = "project_id"
_SK_CLIENT_ID  = "linked_client_id"   # client the conversation is about
_SK_TITLE_SET  = "title_set"
_SK_SAVE_NUDGE = "save_nudge_shown"   # show "save conversation" button only once per session

# ── Load .env ──────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).parent
load_dotenv(BASE_DIR / ".env", override=True)

# ── Paths ──────────────────────────────────────────────────────────────────
MCP_SERVER_DIR = BASE_DIR / "mcp-server"
DB_PATH = MCP_SERVER_DIR / "insurance_broker.db"
OUTPUT_DIR = MCP_SERVER_DIR / "output"
OUTPUT_DIR.mkdir(exist_ok=True)

sys.path.insert(0, str(MCP_SERVER_DIR))

# ── Import broker tools directly ────────────────────────────────────────────
import traceback as _tb
_tool_imports = [
    ("insurance_broker_mcp.tools.client_tools", ["search_clients_fn", "get_client_fn", "create_client_fn", "update_client_fn", "delete_client_fn"]),
    ("insurance_broker_mcp.tools.policy_tools", ["get_renewals_due_fn", "list_policies_fn"]),
    ("insurance_broker_mcp.tools.product_tools", ["search_products_fn", "compare_products_fn"]),
    ("insurance_broker_mcp.tools.offer_tools", ["create_offer_fn", "list_offers_fn"]),
    ("insurance_broker_mcp.tools.claims_tools", ["log_claim_fn", "get_claim_status_fn"]),
    ("insurance_broker_mcp.tools.compliance_tools", ["asf_summary_fn", "bafin_summary_fn", "check_rca_validity_fn"]),
    ("insurance_broker_mcp.tools.email_tools", ["send_offer_email_fn"]),
    ("insurance_broker_mcp.tools.analytics_tools", ["cross_sell_fn"]),
    ("insurance_broker_mcp.tools.calculator_tools", ["calculate_premium_fn", "compare_premiums_live_fn"]),
]
for _mod_name, _names in _tool_imports:
    try:
        _mod = __import__(_mod_name, fromlist=_names)
        for _n in _names:
            globals()[_n] = getattr(_mod, _n)
    except Exception as _e:
        print(f"[app.py] FATAL: Failed to import {_names} from {_mod_name}: {_e}", flush=True)
        _tb.print_exc()
        raise
from insurance_broker_mcp.tools.vehicle_tools import (
    add_vehicle_fn, list_vehicles_fn, search_vehicle_fn,
    get_vehicle_fn, update_vehicle_fn, delete_vehicle_fn,
)
from insurance_broker_mcp.tools.compliance_check_tools import compliance_check_fn
from insurance_broker_mcp.tools.web_tools import check_rca_fn as _playwright_check_rca_fn, browse_web_fn as _playwright_browse_web_fn, scrape_rca_prices_fn as _scrape_rca_prices_fn
from insurance_broker_mcp.tools.drive_tools import (
    upload_to_drive_fn, list_drive_files_fn, get_drive_link_fn,
    sp_upload_fn, sp_list_fn, sp_get_link_fn,
)
from insurance_broker_mcp.tools.rag_tools import (
    broker_search_knowledge_fn,
    broker_upload_document_fn,
    broker_analyze_document_fn,
    broker_kb_status_fn,
    broker_kb_reindex_fn,
)

# ── Safari / mobile login fix ─────────────────────────────────────────────────
# Safari (iOS + macOS WebKit 605.x) sends Content-Type: application/json for
# the /login POST. Chainlit's route uses OAuth2PasswordRequestForm which requires
# application/x-www-form-urlencoded. We replace Chainlit's /login route handler
# with one that accepts BOTH JSON and form-urlencoded bodies.
try:
    from chainlit.server import app as _cl_app_login
    from fastapi import Request as _LoginRequest, Response as _LoginResponse
    import json as _login_json

    async def _safari_compat_login(request: _LoginRequest, response: _LoginResponse):
        """Login handler that accepts both JSON and form-urlencoded bodies.
        Fixes Safari/iOS which sends Content-Type: application/json."""
        import logging as _ll
        _log = _ll.getLogger("login_debug")

        from chainlit.config import config as _cl_config
        from chainlit.server import _authenticate_user as _cl_auth_user

        ct = request.headers.get("content-type", "")
        raw_body = await request.body()

        _log.warning("LOGIN_DEBUG ct=%r body_len=%d body_preview=%r headers=%s",
                     ct, len(raw_body), raw_body[:200],
                     dict(request.headers))

        if "application/json" in ct:
            try:
                data = _login_json.loads(raw_body)
                username = (data.get("username") or data.get("email") or "").strip()
                password = (data.get("password") or "").strip()
            except Exception as _je:
                _log.warning("LOGIN_DEBUG json_parse_error=%s", _je)
                from fastapi import HTTPException as _HE
                raise _HE(status_code=422, detail="Invalid JSON body")
        else:
            # re-construct a receive so form() can read the already-consumed body
            async def _receive():
                return {"type": "http.request", "body": raw_body, "more_body": False}
            from starlette.requests import Request as _SR
            _req2 = _SR(request.scope, _receive)
            form = await _req2.form()
            username = str(form.get("username") or form.get("email") or "").strip()
            password = str(form.get("password") or "").strip()

        _log.warning("LOGIN_DEBUG username=%r password_len=%d", username, len(password))

        if not _cl_config.code.password_auth_callback:
            from fastapi import HTTPException as _HE
            raise _HE(status_code=400, detail="No auth_callback defined")

        user = await _cl_config.code.password_auth_callback(username, password)
        _log.warning("LOGIN_DEBUG auth_result=%r", user)
        return await _cl_auth_user(request, user)

    # Remove Chainlit's /login POST route from the underlying router, then add ours
    _cl_app_login.router.routes = [
        r for r in _cl_app_login.router.routes
        if not (getattr(r, "path", None) == "/login" and
                "POST" in getattr(r, "methods", set()))
    ]
    _cl_app_login.add_api_route("/login", _safari_compat_login, methods=["POST"])

except Exception as _login_fix_err:
    import logging as _lf_log
    _lf_log.getLogger("app").warning("Safari login fix not applied: %s", _login_fix_err)

# ── REST API endpoints for n8n / external automation ─────────────────────────
# Added on chainlit.server.app — only /api/* and /health paths, no Chainlit conflicts
try:
    from chainlit.server import app as _cl_app
    from fastapi import Query as _Query
    from fastapi.responses import JSONResponse as _JSONResponse

    @_cl_app.get("/health")
    async def _health():
        return _JSONResponse({"status": "ok"})

    @_cl_app.get("/api/renewals")
    async def _api_renewals(days: int = _Query(default=45, ge=1, le=365)):
        """Structured renewal list — useful for n8n automation.
        Returns JSON with urgent (≤7 days) and upcoming (≤days) lists.
        """
        import sqlite3 as _sqlite3
        from datetime import date as _date, timedelta as _td
        from pathlib import Path as _Path
        _db = _Path(__file__).parent / "mcp-server" / "insurance_broker.db"
        try:
            _conn = _sqlite3.connect(str(_db))
            _conn.row_factory = _sqlite3.Row
            _today = _date.today().isoformat()
            _cutoff = (_date.today() + _td(days=days)).isoformat()
            _rows = _conn.execute("""
                SELECT p.id, p.client_id, c.name as client_name, c.email as client_email,
                       c.phone as client_phone,
                       p.policy_type, p.insurer, p.policy_number, p.end_date, p.annual_premium, p.currency,
                       CAST(julianday(p.end_date) - julianday('now') AS INTEGER) as days_left
                FROM policies p
                JOIN clients c ON c.id = p.client_id
                WHERE p.status = 'active' AND p.end_date BETWEEN ? AND ?
                ORDER BY p.end_date ASC
            """, (_today, _cutoff)).fetchall()
            _conn.close()
            _items = [dict(r) for r in _rows]
            return _JSONResponse({
                "as_of": _today,
                "days_ahead": days,
                "total": len(_items),
                "urgent": [i for i in _items if i["days_left"] <= 7],
                "upcoming": [i for i in _items if i["days_left"] > 7],
                "all": _items,
            })
        except Exception as _ex:
            return _JSONResponse({"error": str(_ex)}, status_code=500)

    @_cl_app.get("/api/claims/open")
    async def _api_open_claims(max_age_days: int = _Query(default=90, ge=1, le=365)):
        """Return open/investigating claims — useful for n8n follow-up automation."""
        import sqlite3 as _sqlite3
        from pathlib import Path as _Path
        _db = _Path(__file__).parent / "mcp-server" / "insurance_broker.db"
        try:
            _conn = _sqlite3.connect(str(_db))
            _conn.row_factory = _sqlite3.Row
            _rows = _conn.execute("""
                SELECT cl.id as claim_id, cl.client_id, c.name as client_name,
                       c.email as client_email, c.phone as client_phone,
                       cl.policy_id, cl.incident_date, cl.status,
                       cl.damage_estimate, cl.description,
                       cl.insurer_claim_number, cl.notes,
                       CAST(julianday('now') - julianday(cl.incident_date) AS INTEGER) as days_open
                FROM claims cl
                JOIN clients c ON c.id = cl.client_id
                WHERE cl.status IN ('open', 'investigating')
                ORDER BY cl.incident_date ASC
            """).fetchall()
            _conn.close()
            _items = [dict(r) for r in _rows]
            return _JSONResponse({
                "total_open": len(_items),
                "overdue": [i for i in _items if i["days_open"] > max_age_days],
                "claims": _items,
            })
        except Exception as _ex:
            return _JSONResponse({"error": str(_ex)}, status_code=500)

    @_cl_app.get("/api/reports/asf")
    async def _api_asf(month: int = _Query(..., ge=1, le=12), year: int = _Query(..., ge=2020, le=2030)):
        result = asf_summary_fn(month=month, year=year)
        return _JSONResponse({"report": result})

    @_cl_app.get("/api/reports/bafin")
    async def _api_bafin(month: int = _Query(..., ge=1, le=12), year: int = _Query(..., ge=2020, le=2030)):
        result = bafin_summary_fn(month=month, year=year)
        return _JSONResponse({"report": result})

    @_cl_app.get("/api/clients/search")
    async def _api_clients(q: str = _Query(..., min_length=1), limit: int = _Query(default=20, ge=1, le=100)):
        result = search_clients_fn(query=q, limit=limit)
        return _JSONResponse({"clients": result})

    @_cl_app.get("/api/dashboard")
    async def _api_dashboard():
        """Summary dashboard stats — for embedding in external tools."""
        import sqlite3 as _sqlite3
        from datetime import date as _date
        from pathlib import Path as _Path
        _db = _Path(__file__).parent / "mcp-server" / "insurance_broker.db"
        try:
            _conn = _sqlite3.connect(str(_db))
            _conn.row_factory = _sqlite3.Row
            _today = _date.today()
            stats = {
                "as_of": _today.isoformat(),
                "active_policies": _conn.execute("SELECT COUNT(*) FROM policies WHERE status='active'").fetchone()[0],
                "clients": _conn.execute("SELECT COUNT(*) FROM clients").fetchone()[0],
                "open_claims": _conn.execute("SELECT COUNT(*) FROM claims WHERE status IN ('open','investigating')").fetchone()[0],
                "expiring_7":  _conn.execute(
                    "SELECT COUNT(*) FROM policies WHERE status='active' AND julianday(end_date)-julianday('now') BETWEEN 0 AND 7"
                ).fetchone()[0],
                "expiring_30": _conn.execute(
                    "SELECT COUNT(*) FROM policies WHERE status='active' AND julianday(end_date)-julianday('now') BETWEEN 0 AND 30"
                ).fetchone()[0],
                "offers_total": _conn.execute("SELECT COUNT(*) FROM offers").fetchone()[0],
            }
            _conn.close()
            return _JSONResponse(stats)
        except Exception as _ex:
            return _JSONResponse({"error": str(_ex)}, status_code=500)

    @_cl_app.get("/api/debug-from-app")
    async def _debug_from_app():
        """Test if app.py subprocess can reach main.py via localhost."""
        import requests as _rr, asyncio as _aa, os as _os
        port = _os.environ.get("PORT", "8080")
        results = {}
        loop = _aa.get_event_loop()
        def _chk(path):
            try:
                r = _rr.get(f"http://localhost:{port}{path}", timeout=3)
                return {"status": r.status_code, "body": r.text[:100]}
            except Exception as e:
                return {"error": str(e)}
        for path in ["/health", "/cu/status"]:
            results[path] = await loop.run_in_executor(None, lambda p=path: _chk(p))
        return _JSONResponse({"port": port, "pid": __import__("os").getpid(), "results": results})

    print("[INFO] /api/* endpoints mounted on chainlit.server.app")
except Exception as _e:
    print(f"[WARN] API endpoints not mounted: {_e}")

# ── Computer Use shared state ─────────────────────────────────────────────────
# When running via `uvicorn main:app`, the state lives in main.py.
# When running standalone via `chainlit run app.py`, we own the state here.
import asyncio as _asyncio
import uuid as _uuid
from datetime import datetime as _datetime
from collections import defaultdict as _defaultdict

# State for computer-use tasks — shared with main.py via cu_state module
# DO NOT import main here — circular import deadlock (main→chainlit→app→main)
# cu_state is a neutral module with no imports from main/chainlit
from cu_state import _cu_tasks, _cu_results, _cu_agents, _cu_pending

try:
    from chainlit.server import app as _cl_app2
    from fastapi import Request as _Request, Header as _Header
    from fastapi.responses import JSONResponse as _JSONResponse2
    from typing import Optional as _Optional

    @_cl_app2.get("/api/computer-use/tasks")
    async def _cu_get_tasks(
        request: _Request,
        x_agent_id: _Optional[str] = _Header(None),
    ):
        """Local agent polls this to get pending tasks."""
        agent_id = x_agent_id or request.headers.get("X-Agent-ID", "unknown")

        # Return tasks queued for this agent
        pending_ids = _cu_pending.get(agent_id, [])
        tasks = []
        remaining_ids = []
        for tid in pending_ids:
            task = _cu_tasks.get(tid)
            if task and task.get("status") == "pending":
                task["status"] = "dispatched"
                tasks.append(task)
            else:
                remaining_ids.append(tid)
        _cu_pending[agent_id] = remaining_ids

        # Update agent last-seen
        if agent_id in _cu_agents:
            _cu_agents[agent_id]["last_seen"] = _datetime.utcnow().isoformat()

        return _JSONResponse2({"tasks": tasks, "agent_id": agent_id})

    @_cl_app2.post("/api/computer-use/results")
    async def _cu_post_result(request: _Request):
        """Local agent posts task results here."""
        try:
            body = await request.json()
            task_id = body.get("task_id")
            if not task_id:
                return _JSONResponse2({"error": "task_id required"}, status_code=400)
            _cu_results[task_id] = {
                "result": body.get("result", {}),
                "agent_id": body.get("agent_id"),
                "completed_at": body.get("completed_at", _datetime.utcnow().isoformat()),
            }
            if task_id in _cu_tasks:
                _cu_tasks[task_id]["status"] = "completed"
            return _JSONResponse2({"ok": True, "task_id": task_id})
        except Exception as e:
            return _JSONResponse2({"error": str(e)}, status_code=500)

    @_cl_app2.post("/api/computer-use/heartbeat")
    async def _cu_heartbeat(request: _Request):
        """Local agent sends heartbeat to show it's online."""
        try:
            body = await request.json()
            agent_id = body.get("agent_id", "unknown")
            _cu_agents[agent_id] = {
                "agent_id": agent_id,
                "platform": body.get("platform", ""),
                "connectors": body.get("connectors", []),
                "last_seen": _datetime.utcnow().isoformat(),
                "python": body.get("python", ""),
            }
            return _JSONResponse2({"ok": True, "agent_id": agent_id})
        except Exception as e:
            return _JSONResponse2({"error": str(e)}, status_code=500)

    @_cl_app2.get("/api/computer-use/status")
    async def _cu_status():
        """Returns all online agents and their capabilities."""
        from datetime import datetime, timezone, timedelta
        online = []
        for agent_id, info in _cu_agents.items():
            try:
                last = _datetime.fromisoformat(info["last_seen"])
                delta = (_datetime.utcnow() - last).total_seconds()
                if delta < 120:  # online if heartbeat within 2 min
                    online.append({**info, "online": True, "seconds_ago": round(delta)})
            except Exception:
                pass
        return _JSONResponse2({
            "agents_online": len(online),
            "agents": online,
            "total_tasks": len(_cu_tasks),
            "pending_results": sum(1 for t in _cu_tasks.values() if t.get("status") == "pending"),
        })

    print("[INFO] /api/computer-use/* endpoints mounted")
except Exception as _e2:
    print(f"[WARN] Computer Use API endpoints not mounted: {_e2}")


def _cu_get_base_url() -> str:
    """Get the base URL for inter-process HTTP calls.

    app.py and main.py run in the SAME uvicorn process (mount_chainlit is in-process).
    We can use localhost directly — much faster and more reliable than self-calls
    via the public Cloud Run URL (which adds ~200ms RTT and can timeout).

    Chainlit intercepts /api/* routes on localhost, but our /cu/* routes are
    mounted on FastAPI in main.py BEFORE mount_chainlit, so they work fine.
    """
    port = os.environ.get("PORT", "8080")
    return f"http://localhost:{port}"


def _cu_enqueue_task(connector: str, action: str, params: dict,
                      agent_id: str = "default", timeout: int = 120,
                      credentials: dict = None) -> str:
    """
    Enqueue a computer-use task via HTTP to the FastAPI server (main.py).
    Uses the public URL because Chainlit intercepts localhost requests.
    """
    import requests as _req
    task_id = str(_uuid.uuid4())
    task = {
        "task_id": task_id,
        "connector": connector,
        "action": action,
        "params": params,
        "credentials": credentials or {},
        "timeout": timeout,
        "status": "pending",
        "created_at": _datetime.utcnow().isoformat(),
    }
    base = _cu_get_base_url()
    try:
        _req.post(
            f"{base}/cu/enqueue",
            json={"task": task, "agent_id": agent_id},
            timeout=10,
        )
    except Exception:
        # Fallback: write directly to cu_state (works if same process)
        _cu_tasks[task_id] = task
        _cu_pending[agent_id].append(task_id)
    return task_id


async def _cu_wait_result(task_id: str, timeout: int = 130) -> dict:
    """Wait for a task result by polling /cu/result via the public URL."""
    import requests as _req
    base = _cu_get_base_url()
    url = f"{base}/cu/result/{task_id}"
    loop = _asyncio.get_event_loop()
    start = loop.time()

    def _poll():
        try:
            r = _req.get(url, timeout=5)
            return r.json()
        except Exception:
            return {}

    while (loop.time() - start) < timeout:
        data = await loop.run_in_executor(None, _poll)
        if data.get("ready"):
            return data.get("result", {})
        await _asyncio.sleep(2)
    return {"success": False, "error": f"Task {task_id} timed out after {timeout}s — is the local agent running?"}

# ── Anthropic client ──────────────────────────────────────────────────────────
_anthropic_key = os.environ.get("ANTHROPIC_API_KEY")
if not _anthropic_key:
    raise RuntimeError("ANTHROPIC_API_KEY not set. Add it to .env file in the project root.")

client = anthropic.Anthropic(api_key=_anthropic_key)
MODEL = os.environ.get("CLAUDE_MODEL", "claude-sonnet-4-5")

# ── RAG Context Injection ─────────────────────────────────────────────────────
def _get_rag_context(user_message: str) -> str:
    """Retrieve relevant knowledge and inject into system prompt.
    Self-contained — no imports from main.py to avoid circular imports."""
    try:
        import sqlite3 as _sq_rag
        _db_path = str(_Path(__file__).parent / "mcp-server" / "insurance_broker.db")
        conn = _sq_rag.connect(_db_path)
        conn.row_factory = _sq_rag.Row
        words = [w.strip() for w in user_message.lower().split() if len(w.strip()) > 2]
        if not words:
            conn.close()
            return ""
        where_clauses = " OR ".join(["LOWER(content) LIKE ?"] * len(words))
        params = [f"%{w}%" for w in words]
        rows = conn.execute(
            f"SELECT id, hook, content FROM alex_knowledge "
            f"WHERE ({where_clauses}) "
            f"ORDER BY relevance_score DESC, created_at DESC LIMIT 5",
            params
        ).fetchall()
        if not rows:
            conn.close()
            return ""
        # Bump times_used
        ids = [r["id"] for r in rows]
        placeholders = ",".join(["?"] * len(ids))
        conn.execute(
            f"UPDATE alex_knowledge SET times_used = times_used + 1, "
            f"last_used_at = datetime('now') WHERE id IN ({placeholders})", ids
        )
        conn.commit()
        conn.close()
        context = "\n\n## What you've learned from previous interactions:\n"
        for r in rows:
            context += f"- [{r['hook']}] {r['content']}\n"
        context += "\nUse this knowledge to give better, more personalized responses.\n"
        return context
    except Exception:
        return ""


# ── System prompt ─────────────────────────────────────────────────────────────
SYSTEM_PROMPT = """You are Alex, a friendly and proactive AI assistant for an insurance brokerage (ASF Romania / BaFin Germany).

You talk naturally with brokers in whatever language they use (Romanian, English, German). You NEVER return raw errors — if something fails, you explain simply and offer alternatives.

## Your Personality
- Warm, concise, professional
- Always suggest the next logical step
- When unsure what the broker wants, ask one clear question
- Never say "I cannot do that" — always find a way or offer an alternative
- Respond in the same language the broker writes in

## Tools Available
- broker_search_clients — caută client după nume, telefon, email
- broker_get_client — profil complet cu polițe
- broker_create_client — adaugă client nou
- broker_update_client — corectează datele unui client existent (nume, telefon, email, etc.)
- broker_delete_client — șterge client (refuză dacă are polițe active)
- broker_search_products — caută produse (RCA, CASCO, PAD, CMR, KFZ, LIFE, HEALTH)
- broker_compare_products — comparație tabelară
- broker_create_offer — generează ofertă profesională (PDF + XLSX + DOCX)
- broker_list_offers — listează ofertele generate
- broker_send_offer_email — trimite oferta pe email clientului
- broker_get_renewals_due — polițe care expiră curând
- broker_list_policies — listează polițe
- broker_log_claim — înregistrează daună
- broker_get_claim_status — status dosar daună
- broker_asf_summary — raport lunar ASF (Romania)
- broker_bafin_summary — raport lunar BaFin (Germania)
- broker_check_rca_validity — verifică valabilitate RCA
- broker_cross_sell — analizează portofoliul clientului și sugerează produse lipsă
- broker_compare_premiums_live — compară prețurile RCA/CASCO de la TOȚI asigurătorii (Allianz, Generali, Omniasig, Groupama, Uniqa, Asirom, Euroins, Grawe). Acceptă: vehicle_type (autoturism/SUV/camion/motocicleta/autoutilitara), vehicle_year (an fabricație), fuel_type (benzina/diesel/GPL/hybrid/electric), period_months (6 sau 12). Returnează tabel sortat cel mai ieftin primul. Folosește oricând brokerul întreabă despre prețuri RCA/CASCO. Tarife orientative 2024-2025 ASF.
- broker_add_vehicle — adaugă vehicul la un client (obligatoriu: client_id + plate_number; opțional: make, model, year, engine_cc, fuel_type, vehicle_type, vehicle_value, vin)
- broker_list_vehicles — listează vehiculele unui client cu polițe asociate
- broker_search_vehicle — caută vehicul după nr. înmatriculare, VIN, marcă sau model
- broker_get_vehicle — detalii complete vehicul cu toate polițele
- broker_update_vehicle — actualizează datele unui vehicul
- broker_delete_vehicle — șterge vehicul (refuză dacă are polițe active)
- broker_compliance_check — verifică completitudinea dosarului client (documente, polițe, conformitate)
- broker_save_conversation — salvează și asociază conversația curentă cu un client (pentru istoric). Apelează când brokerul spune "salvează", "linkuiește la Ionescu", etc.
- broker_check_rca — verifică RCA în timp real pe portalul AIDA/BAAR via browser headless pe server (NU necesită agent local). Returnează: rca_valid, expiry_date, insurer, policy_number, coverage_type, insured_sum, days_until_expiry, captcha_blocked, from_cache (cache TTL 6h), screenshot_b64 (la eșec).
- broker_scrape_rca_prices — **prețuri RCA REALE** via pint.ro (agregator): introduci numărul de înmatriculare, returnează oferte reale de la asigurători, sortat cel mai ieftin primul. Folosește când brokerul are numărul de înmatriculare și vrea prețuri reale. Dacă eșuează, fallback automat la broker_compare_premiums_live.
- broker_browse_web — accesează orice URL public și extrage text sau tabele (NU necesită agent local)
- broker_computer_use_status — verifică dacă agentul local e conectat (necesar doar pentru desktop apps sau intranet)
- broker_run_task — execută task pe calculatorul angajatului. Connectors disponibili: `desktop_generic` (desktop apps, rețea internă), `cedam` (verificare RCA via portal ASF/CEDAM), `web_generic` (orice site web via browser local), `anthropic_computer_use` (computer use avansat cu Claude Vision).
- broker_drive_upload — încarcă un fișier generat (ofertă PDF, raport XLSX/DOCX) în Google Drive; returnează link partajabil
- broker_drive_list — listează fișierele din folderul Google Drive al brokerului
- broker_drive_get_link — obține linkul unui fișier deja încărcat în Google Drive
- broker_sharepoint_upload — încarcă un fișier în SharePoint (Microsoft 365) via Microsoft Graph; returnează link intern
- broker_sharepoint_list — listează fișierele din folderul SharePoint configurat
- broker_sharepoint_get_link — obține linkul unui fișier deja încărcat în SharePoint
- broker_search_knowledge — căutare semantică în knowledge base (RAG): produse, ghiduri daune, conformitate, FAQ. Folosește pentru întrebări vagi/naturale despre acoperiri, excluderi, proceduri.
- broker_upload_document — încarcă PDF/imagine în Files API Anthropic pentru analiză persistentă. Returnează file_id reutilizabil.
- broker_analyze_document — **[BETA — în testare]** analizează document cu Claude Vision: polițe scanate, facturi, constatare amiabilă, buletin. Acceptă path local SAU file_id. IMPORTANT: rezultatele pot fi inexacte — verificați manual datele extrase înainte de a le folosi. Analiza fotografiilor de daune (doc_type="claim_photo") este dezactivată în producție — disponibilă în Faza 2 după validare RAG.
- broker_kb_status — starea knowledge base (număr chunks indexate, categorii)
- broker_kb_reindex — re-indexează knowledge base (după adăugare produse noi sau actualizare docs/)
- broker_list_output_files — listează ofertele/rapoartele generate și salvate local (PDF, TXT, XLSX, DOCX). Folosește când brokerul întreabă "ce oferte am generat?", "curăță fișierele vechi", "arată-mi ce am salvat".
- broker_send_claim_questionnaire — trimite chestionar de daune unui client pentru un claim specific. Auto-selectează template KFZ-Schaden sau Haftpflicht pe baza tipului de poliță. Trimite prin email + WhatsApp.
- broker_auto_send_questionnaires — detectează TOATE claimurile deschise fără chestionar și trimite automat formularul potrivit (KFZ-Schaden sau Haftpflicht) fiecărui client. Returnează câte formulare au fost trimise.
- broker_check_form_status — rezumat statusul tuturor chestionarelor: câte trimise, în curs, completate, restante. Include detalii client și referință.
- broker_run_form_reminders — rulează follow-up: verifică toate chestionarele incomplete și trimite remindere clienților care nu au completat după 4, 8 sau 12 zile.

## Când să folosești ce tool:
- **RCA verificare** → `broker_check_rca` (instant, fără agent). Dacă răspunsul conține `captcha_blocked: true` → folosește automat `broker_run_task` cu connector=`cedam` (agent local, browser vizibil, ocolește CAPTCHA).
- **Orice site web public** → `broker_browse_web` (fără agent)
- **Deschide aplicație + scrie text** (TextEdit, Notes, Word, Excel) → `broker_run_task` cu connector=`desktop_generic`, action=`open_app_and_type`, params={app: "TextEdit", text: "textul exact"} — OBLIGATORIU această acțiune, NU run_task!
- **Calculator** → `broker_run_task` cu connector=`desktop_generic`, action=`run_task`, params.instruction="deschide Calculator si calculeaza 2+2"
- **Rețea internă** (intranet, VPN) → `broker_run_task` cu connector=`desktop_generic`, action=`run_task`
- **Verificare RCA via agent local** (dacă `broker_check_rca` nu merge) → `broker_run_task` cu connector=`cedam`, action=`check_rca`, params={plate: "B123ABC"}
- **Automatizare site web via browser local** → `broker_run_task` cu connector=`web_generic`, action=`run_task`
- **Computer use avansat cu AI Vision** → `broker_run_task` cu connector=`anthropic_computer_use`, action=`run_task`
- **Salvează ofertă/raport în Google Drive** → `broker_drive_upload` cu filename=numele fișierului din output/; apoi dă linkul clientului
- **Salvează în SharePoint (Microsoft 365)** → `broker_sharepoint_upload` — pentru firme pe M365
- **Listează fișierele salvate** → `broker_drive_list` sau `broker_sharepoint_list`
- **Întrebare vagă despre acoperiri/excluderi/proceduri** (ex: "ce acoperă CASCO?", "ce documente trebuie la daune Allianz?") → `broker_search_knowledge` ÎNAINTE de `broker_search_products`
- **Broker uploadează un PDF/imagine** → `broker_upload_document` → obții file_id → `broker_analyze_document(file_id)` pentru extragere date (BETA — verificați rezultatele manual)
- **Analizează poliță scanată** → `broker_analyze_document(doc_type="policy")` → extrage număr poliță, date, primă, acoperire (BETA)
- **Analizează poze daune** → doc_type="claim_photo" este DEZACTIVAT în producție. Spune brokerului că această funcție este în testare și va fi disponibilă în Faza 2.
- **Procesează constatare amiabilă** → `broker_analyze_document(doc_type="constatare")` → extrage ambii șoferi (BETA — verificați manual)
- **Comparator prețuri RCA/CASCO** (orice variantă: "compară", "cât costă", "care e mai ieftin", "vreau să văd prețurile") → `broker_compare_premiums_live` → tabel complet cu toți asigurătorii sortat după preț. IMPORTANT: trimite și vehicle_type, vehicle_year, fuel_type, period_months dacă le cunoști!
- **Client spune nr. mașinii** → `broker_search_vehicle(query=nr_inmatriculare)` → dacă nu există, creează cu `broker_add_vehicle`
- **RCA/CASCO pentru un client** → `broker_list_vehicles(client_id)` → ia datele vehiculului → `broker_compare_premiums_live` cu engine_cc, vehicle_type, vehicle_year, fuel_type din vehicul
- **Onboarding client cu mașină** → `broker_create_client` → `broker_add_vehicle` → `broker_compare_premiums_live` → `broker_create_offer` — flow complet
- **Generare ofertă auto** → include în ofertă: marca, modelul, anul, nr. înmatriculare vehiculului
- **Trimite chestionar la un claim** → `broker_send_claim_questionnaire(claim_id)` — auto-selectează KFZ sau Haftpflicht
- **Trimite chestionare la TOATE claimurile fără formular** → `broker_auto_send_questionnaires` — detectează claimuri open/investigating
- **Verifică status chestionare** → `broker_check_form_status` — afișează total, status, restante
- **Trimite remindere chestionare** → `broker_run_form_reminders` — 4/8/12 zile, max 3 remindere
- **Flow complet daune** → `broker_log_claim` → `broker_send_claim_questionnaire` → (client completează) → `broker_check_form_status`
- INTERZIS: action `fill_form` pentru sarcini desktop simple — folosește `run_task` cu instrucțiune naturală.
- INTERZIS: action `run_task` când utilizatorul cere să deschizi o aplicație și să scrii text — folosește `open_app_and_type`.

## CRITICAL: Always use tools — NEVER answer from memory

**You have NO knowledge of real products, clients, or policies.** All data lives in the database.
- NEVER list insurance products, prices, or insurers from your training — you don't have this data
- NEVER say "Here are the top HEALTH products: NN, Allianz..." without calling broker_search_products first
- NEVER confirm or recommend a product without calling the tool first
- If asked about products → call broker_search_products IMMEDIATELY, then show the real results
- If asked about a client → call broker_search_clients or broker_get_client first
- **NEVER say price comparison is unavailable** — `broker_compare_premiums_live` is ALWAYS available. Call it immediately for ANY RCA/CASCO price question.

**Rule:** If you are about to mention a product name, insurer, price, or policy — STOP and call the tool instead.

## How to Handle Any Request
1. **Understand intent** — even vague requests ("fa ceva cu asta", "vreau oferta", "ce mai am de facut")
2. **Always call a tool first** — never describe data you haven't fetched yet
3. **Chain actions** — search client → create if missing → add vehicle if auto insurance → compare premiums → create offer, all in one flow when the broker says "onboard" or "fa oferta"
4. **After offer is generated** — it's already shown in chat. Don't regenerate. Ask if they want to email it or modify something
5. **If a tool returns no results** — suggest alternatives: "Nu am găsit produse HEALTH, dar am LIFE cu acoperire similară. Generez oferta cu astea?"
6. **If something fails** — say what happened in plain language and offer 2-3 options to continue

## Document Upload Workflow
When a doc is analyzed and broker confirms data is correct:
1. Search DB using CLIENT name (never clinic/issuer phone or name) → create client if not found
2. Identify best product type from document context
3. Ask ONE question: "Am creat clientul X. Generez ofertă de HEALTH?" → on any affirmative → do it immediately

CRITICAL for document OCR:
- The CLIENT is the person the document is ABOUT (patient, vehicle owner, policy holder)
- The ISSUER is the clinic/insurer/company that issued it — NEVER use issuer phone/name as client data
- If broker says "corecție" / "greșit" / "numele e..." → update your understanding and use the corrected data
- NEVER search using a clinic phone number or company name as if it were a client

## Offer Workflow
- "da" / "yes" / "fă" / "generează" / "make it" after seeing products = call broker_create_offer IMMEDIATELY, no more questions
- After offer is generated, the UI will ask the broker to approve/edit/download — you don't need to ask again
- If broker says "aprobă" / "trimite" / "send" after seeing offer → call broker_send_offer_email right away
- If broker says "editează" / "modifică" + describes changes → call broker_create_offer again with updated parameters
- Never call broker_create_offer again just to "show" an offer — use the session cache instead
- NEVER get stuck waiting — always suggest the next logical action

## Language & Format
- Match the broker's language automatically
- Keep responses short — max 3-4 lines unless showing data tables
- Use ✅ ⚠️ 📄 💡 sparingly for clarity
- Amounts: RON for RO products, EUR for DE/KFZ

## Compliance
Never show full CNP/passport numbers — use client IDs only."""

# ── Tool definitions for Anthropic function calling ───────────────────────────
TOOLS = [
    {
        "name": "broker_search_clients",
        "description": "Search clients by name, phone, or email address.",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Name, phone, or email to search"},
                "limit": {"type": "integer", "description": "Max results (default 10)"},
            },
            "required": ["query"],
        },
    },
    {
        "name": "broker_get_client",
        "description": "Get full client profile including all their policies.",
        "input_schema": {
            "type": "object",
            "properties": {
                "client_id": {"type": "string", "description": "Client ID (e.g. CLI001)"},
            },
            "required": ["client_id"],
        },
    },
    {
        "name": "broker_create_client",
        "description": "Create a new client in the database.",
        "input_schema": {
            "type": "object",
            "properties": {
                "name": {"type": "string"},
                "phone": {"type": "string"},
                "email": {"type": "string"},
                "address": {"type": "string"},
                "client_type": {"type": "string", "description": "individual or company"},
                "country": {"type": "string", "description": "RO or DE"},
                "source": {"type": "string"},
            },
            "required": ["name", "phone"],
        },
    },
    {
        "name": "broker_update_client",
        "description": "Update an existing client's details (name, phone, email, address, etc.). Only provided fields are changed.",
        "input_schema": {
            "type": "object",
            "properties": {
                "client_id": {"type": "string", "description": "Client ID from broker_search_clients"},
                "name": {"type": "string"},
                "phone": {"type": "string"},
                "email": {"type": "string"},
                "address": {"type": "string"},
                "id_number": {"type": "string"},
                "country": {"type": "string", "description": "RO or DE"},
                "client_type": {"type": "string", "description": "individual or company"},
                "notes": {"type": "string"},
            },
            "required": ["client_id"],
        },
    },
    {
        "name": "broker_delete_client",
        "description": "Delete a client. Will refuse if they have active policies. Use broker_search_clients first to get client_id.",
        "input_schema": {
            "type": "object",
            "properties": {
                "client_id": {"type": "string", "description": "Client ID to delete"},
            },
            "required": ["client_id"],
        },
    },
    {
        "name": "broker_search_products",
        "description": "Search available insurance products from all partner insurers.",
        "input_schema": {
            "type": "object",
            "properties": {
                "product_type": {"type": "string", "description": "Type: RCA, CASCO, PAD, HOME, LIFE, CMR, KFZ, LIABILITY, TRANSPORT, HEALTH"},
                "country": {"type": "string", "description": "RO or DE"},
            },
            "required": ["product_type"],
        },
    },
    {
        "name": "broker_compare_products",
        "description": "Compare multiple insurance products side by side.",
        "input_schema": {
            "type": "object",
            "properties": {
                "product_ids": {"type": "string", "description": "Comma-separated product IDs from broker_search_products"},
            },
            "required": ["product_ids"],
        },
    },
    {
        "name": "broker_create_offer",
        "description": "Generate a professional insurance offer document for a client.",
        "input_schema": {
            "type": "object",
            "properties": {
                "client_id": {"type": "string"},
                "product_ids": {"type": "string", "description": "Comma-separated product IDs"},
                "language": {"type": "string", "description": "en, ro, or de"},
                "valid_days": {"type": "integer"},
                "notes": {"type": "string"},
                "format": {"type": "string", "description": "text or pdf"},
            },
            "required": ["client_id", "product_ids"],
        },
    },
    {
        "name": "broker_list_offers",
        "description": "List generated offers, optionally filtered by client or status.",
        "input_schema": {
            "type": "object",
            "properties": {
                "client_id": {"type": "string"},
                "status": {"type": "string", "description": "sent, accepted, or expired"},
            },
        },
    },
    {
        "name": "broker_send_offer_email",
        "description": "Send a generated offer by email to the client. Fetches client email from DB if to_email not provided.",
        "input_schema": {
            "type": "object",
            "properties": {
                "offer_id": {"type": "string", "description": "Offer ID from broker_list_offers"},
                "to_email": {"type": "string", "description": "Recipient email (optional, uses client email from DB if not set)"},
                "subject": {"type": "string", "description": "Email subject (optional)"},
            },
            "required": ["offer_id"],
        },
    },
    {
        "name": "broker_get_renewals_due",
        "description": "Get policies expiring within the next N days.",
        "input_schema": {
            "type": "object",
            "properties": {
                "days_ahead": {"type": "integer", "description": "Number of days to look ahead (default 30)"},
            },
        },
    },
    {
        "name": "broker_list_policies",
        "description": "List policies, optionally filtered by client ID or status.",
        "input_schema": {
            "type": "object",
            "properties": {
                "client_id": {"type": "string"},
                "status": {"type": "string", "description": "active, expired, or cancelled"},
            },
        },
    },
    {
        "name": "broker_log_claim",
        "description": "Register a new damage/claims report for a client.",
        "input_schema": {
            "type": "object",
            "properties": {
                "client_id": {"type": "string"},
                "policy_id": {"type": "string"},
                "incident_date": {"type": "string", "description": "YYYY-MM-DD"},
                "description": {"type": "string"},
                "damage_estimate": {"type": "number"},
            },
            "required": ["client_id", "policy_id", "incident_date", "description"],
        },
    },
    {
        "name": "broker_get_claim_status",
        "description": "Get the status of an existing claim.",
        "input_schema": {
            "type": "object",
            "properties": {
                "claim_id": {"type": "string"},
            },
            "required": ["claim_id"],
        },
    },
    {
        "name": "broker_asf_summary",
        "description": "Generate monthly ASF (Romanian Financial Supervisory Authority) report.",
        "input_schema": {
            "type": "object",
            "properties": {
                "month": {"type": "integer", "description": "Month number 1-12"},
                "year": {"type": "integer"},
            },
            "required": ["month", "year"],
        },
    },
    {
        "name": "broker_bafin_summary",
        "description": "Generate monthly BaFin (German) regulatory report.",
        "input_schema": {
            "type": "object",
            "properties": {
                "month": {"type": "integer", "description": "Month number 1-12"},
                "year": {"type": "integer"},
            },
            "required": ["month", "year"],
        },
    },
    {
        "name": "broker_check_rca_validity",
        "description": "Check RCA (mandatory Romanian motor insurance) validity for a client.",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Client name or policy number"},
            },
            "required": ["query"],
        },
    },
    {
        "name": "broker_cross_sell",
        "description": "Analyze a client's portfolio and suggest missing insurance products for cross-selling opportunities.",
        "input_schema": {
            "type": "object",
            "properties": {
                "client_id": {"type": "string", "description": "Client ID (e.g. CLI001)"},
            },
            "required": ["client_id"],
        },
    },
    {
        "name": "broker_compare_premiums_live",
        "description": (
            "PRIMARY tool for ANY price comparison request. "
            "Compare RCA or CASCO prices from ALL 8 major Romanian insurers simultaneously "
            "(Allianz-Tiriac, Generali, Omniasig, Groupama, Uniqa, Asirom, Euroins, Grawe). "
            "Returns a ranked comparison table sorted cheapest first, with annual premium, monthly cost, insurer rating, and price difference. "
            "Accepts vehicle details: vehicle_type (autoturism/SUV/camion/motocicleta/autoutilitara), vehicle_year, "
            "fuel_type (benzina/diesel/GPL/hybrid/electric), period_months (6 or 12). "
            "Use this for: 'compară prețurile', 'cât costă RCA', 'care e cel mai ieftin', 'vreau să văd toate prețurile', "
            "'compara asiguratorii', 'compare prices', 'cheapest RCA', 'price comparison'. "
            "Sources: market rates 2024-2025 based on public ASF data, updated quarterly."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "product_type": {"type": "string", "description": "RCA or CASCO"},
                "age": {"type": "integer", "description": "Driver age (default 35)"},
                "engine_cc": {"type": "integer", "description": "Engine capacity in cc (default 1600)"},
                "bonus_malus_class": {"type": "string", "description": "Bonus-Malus class B0-B14 or M1-M8 (default B0)"},
                "zone": {"type": "string", "description": "City or Urban/Rural — e.g. Bucuresti, Cluj, Urban, Rural (default Urban)"},
                "vehicle_value": {"type": "number", "description": "Vehicle value in RON — required for CASCO"},
                "country": {"type": "string", "description": "RO (default)"},
                "insurers": {"type": "string", "description": "Optional: comma-separated list of specific insurers to compare (e.g. 'Allianz-Tiriac,Groupama'). Leave empty for all."},
                "vehicle_type": {"type": "string", "description": "autoturism, SUV, camion, motocicleta, autoutilitara (default autoturism)"},
                "vehicle_year": {"type": "integer", "description": "Vehicle manufacturing year (e.g. 2021). Affects age factor."},
                "fuel_type": {"type": "string", "description": "benzina, diesel, GPL, hybrid, electric (default benzina)"},
                "period_months": {"type": "integer", "description": "Insurance period: 6 or 12 months (default 12)"},
            },
            "required": ["product_type"],
        },
    },
    # ── Vehicle management tools ──────────────────────────────────────────────
    {
        "name": "broker_add_vehicle",
        "description": "Add a vehicle to a client. Required: client_id, plate_number. Optional: make, model, year, engine_cc, fuel_type, vehicle_type, vehicle_value, vin.",
        "input_schema": {
            "type": "object",
            "properties": {
                "client_id": {"type": "string", "description": "Client ID (e.g. CLI001)"},
                "plate_number": {"type": "string", "description": "Nr. inmatriculare (e.g. B123ABC, CJ01XYZ, F-AB-1234)"},
                "make": {"type": "string", "description": "Marca vehiculului (e.g. Dacia, BMW, VW, Mercedes)"},
                "model": {"type": "string", "description": "Modelul (e.g. Duster, X5, Golf, Actros)"},
                "year": {"type": "integer", "description": "Anul fabricatiei (e.g. 2021)"},
                "engine_cc": {"type": "integer", "description": "Capacitate cilindrica (e.g. 1461)"},
                "engine_power_kw": {"type": "integer", "description": "Putere motor kW (e.g. 84)"},
                "fuel_type": {"type": "string", "description": "benzina, diesel, GPL, hybrid, electric (default benzina)"},
                "vehicle_type": {"type": "string", "description": "autoturism, SUV, camion, motocicleta, autoutilitara (default autoturism)"},
                "color": {"type": "string", "description": "Culoare vehicul"},
                "vehicle_value": {"type": "number", "description": "Valoare vehicul in RON (necesar pt CASCO)"},
                "vin": {"type": "string", "description": "Serie sasiu / VIN (17 caractere)"},
                "seats": {"type": "integer", "description": "Numar locuri (default 5)"},
                "gross_weight_kg": {"type": "integer", "description": "Masa totala maxima kg"},
                "registration_date": {"type": "string", "description": "Data primei inmatriculari YYYY-MM-DD"},
                "notes": {"type": "string", "description": "Note suplimentare"},
            },
            "required": ["client_id", "plate_number"],
        },
    },
    {
        "name": "broker_list_vehicles",
        "description": "List all vehicles for a client, with active policy associations.",
        "input_schema": {
            "type": "object",
            "properties": {
                "client_id": {"type": "string", "description": "Client ID (e.g. CLI001)"},
            },
            "required": ["client_id"],
        },
    },
    {
        "name": "broker_search_vehicle",
        "description": "Search vehicles by plate number, VIN, make, or model across all clients.",
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Nr. inmatriculare, VIN, marca, sau model"},
            },
            "required": ["query"],
        },
    },
    {
        "name": "broker_get_vehicle",
        "description": "Get full vehicle details with all associated policies (active and expired).",
        "input_schema": {
            "type": "object",
            "properties": {
                "vehicle_id": {"type": "string", "description": "Vehicle ID (e.g. VEH001)"},
            },
            "required": ["vehicle_id"],
        },
    },
    {
        "name": "broker_update_vehicle",
        "description": "Update a vehicle's details. Only provided fields are changed.",
        "input_schema": {
            "type": "object",
            "properties": {
                "vehicle_id": {"type": "string", "description": "Vehicle ID (e.g. VEH001)"},
                "plate_number": {"type": "string"},
                "make": {"type": "string"},
                "model": {"type": "string"},
                "year": {"type": "integer"},
                "engine_cc": {"type": "integer"},
                "engine_power_kw": {"type": "integer"},
                "fuel_type": {"type": "string"},
                "vehicle_type": {"type": "string"},
                "color": {"type": "string"},
                "vehicle_value": {"type": "number"},
                "vin": {"type": "string"},
                "notes": {"type": "string"},
            },
            "required": ["vehicle_id"],
        },
    },
    {
        "name": "broker_delete_vehicle",
        "description": "Delete a vehicle. Refuses if it has active policies.",
        "input_schema": {
            "type": "object",
            "properties": {
                "vehicle_id": {"type": "string", "description": "Vehicle ID to delete"},
            },
            "required": ["vehicle_id"],
        },
    },
    {
        "name": "broker_compliance_check",
        "description": "Check compliance status for a client — missing documents, expiring policies, mandatory product gaps, regulatory issues. Returns a compliance score 0-100.",
        "input_schema": {
            "type": "object",
            "properties": {
                "client_id": {"type": "string", "description": "Client ID (e.g. CLI001)"},
            },
            "required": ["client_id"],
        },
    },
    {
        "name": "broker_save_conversation",
        "description": (
            "Save and link the current conversation to a specific client. "
            "Call this when the broker says something like 'save this conversation', "
            "'link to Ionescu', 'save chat about CLI001', etc. "
            "After calling this tool, the conversation will appear under that client's "
            "history in the '📁 Conversation history by client' section. "
            "Use broker_search_clients first to find the client_id if you don't have it."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "client_id": {
                    "type": "string",
                    "description": "Client ID to link this conversation to (e.g. CLI001)",
                },
                "title": {
                    "type": "string",
                    "description": "Optional title for this conversation (max 80 chars). If omitted, keeps current title.",
                },
            },
            "required": ["client_id"],
        },
    },
    # ── Web automation tools (Playwright on Cloud Run — no local agent needed) ──
    {
        "name": "broker_check_rca",
        "description": (
            "Verifică valabilitatea poliței RCA pentru un număr de înmatriculare, "
            "accesând portalul AIDA/BAAR în timp real via browser headless pe server. "
            "Nu necesită agent local — rulează direct pe Cloud Run. Cache TTL 6h — "
            "dacă plăcuța a fost verificată recent, răspuns instant fără request nou. "
            "Returnează: rca_valid (bool), expiry_date, insurer, policy_number, "
            "coverage_type (RCA/CASCO/CMR etc.), insured_sum, days_until_expiry, "
            "captcha_blocked (bool — dacă True, folosește broker_run_task cu connector=cedam), "
            "from_cache (bool), cached_at, screenshot_b64 (base64 PNG la eșec, pentru debugging)."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "plate": {"type": "string", "description": "Numărul de înmatriculare (ex: B123ABC, B 123 ABC, CJ12XYZ)"},
            },
            "required": ["plate"],
        },
    },
    {
        "name": "broker_browse_web",
        "description": (
            "Accesează un URL public și extrage conținutul paginii (text sau tabele). "
            "Util pentru verificări pe site-uri externe: prețuri, știri, portaluri publice, etc. "
            "Nu necesită agent local — rulează headless pe Cloud Run."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "URL-ul de accesat (trebuie să fie public, fără autentificare)"},
                "query": {"type": "string", "description": "Ce anume căutăm/extragem din pagină (pentru context)"},
                "extract_type": {"type": "string", "description": "'text' (default) pentru text complet, 'table' pentru tabele HTML"},
            },
            "required": ["url"],
        },
    },
    {
        "name": "broker_scrape_rca_prices",
        "description": (
            "Obține prețuri RCA REALE de la asigurători via pint.ro (agregator de asigurări). "
            "Input: numărul de înmatriculare al vehiculului. "
            "pint.ro auto-completează datele vehiculului din numărul de înmatriculare și returnează "
            "oferte reale de la asigurătorii parteneri. "
            "Returnează tabel comparativ sortat cel mai ieftin primul, cu preț anual și lunar. "
            "Cache TTL 2h. Folosește când brokerul are numărul de înmatriculare și vrea prețuri reale (nu estimări). "
            "Fallback automat la broker_compare_premiums_live dacă scraping-ul eșuează."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "plate": {"type": "string", "description": "Numărul de înmatriculare (ex: B123ABC, CJ12XYZ, B 123 ABC)"},
            },
            "required": ["plate"],
        },
    },
    # ── Computer Use tools (local agent — for desktop apps and intranets) ──
    {
        "name": "broker_computer_use_status",
        "description": "Check if the local computer-use agent is online on the employee's machine. Returns list of connected agents and available connectors (cedam, web_generic, desktop_generic, etc.).",
        "input_schema": {
            "type": "object",
            "properties": {},
        },
    },
    {
        "name": "broker_run_task",
        "description": (
            "Run a computer-use task on the employee's local machine. "
            "The local agent executes the task (browser automation or desktop automation) and returns the result. "
            "Use connector='cedam' for RCA checks, connector='web_generic' for any website, connector='desktop_generic' for Windows/Mac apps. "
            "Actions: extract, fill_form, navigate, click, screenshot, check_rca, read_screen, run_task, open_app_and_type. "
            "IMPORTANT: For opening a desktop app and typing text, ALWAYS use action='open_app_and_type' with params={app: 'TextEdit', text: 'exact text here'} — NOT run_task."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "connector": {"type": "string", "description": "Connector to use: cedam, web_generic, desktop_generic"},
                "action": {"type": "string", "description": "Action: extract, fill_form, navigate, click, screenshot, check_rca, read_screen, run_task, login, open_app_and_type"},
                "params": {
                    "type": "object",
                    "description": "Action parameters. For open_app_and_type: {app: 'TextEdit', text: 'exact text to type'}. For run_task: {instruction: '...'}. For check_rca: {plate: 'B123ABC'}.",
                    "properties": {
                        "query": {"type": "string"},
                        "url": {"type": "string"},
                        "plate": {"type": "string", "description": "License plate for RCA check"},
                        "instruction": {"type": "string", "description": "Natural language instruction for run_task (desktop apps)"},
                        "max_steps": {"type": "integer", "description": "Max automation steps (default 10)"},
                        "app": {"type": "string", "description": "App name for open_app_and_type, e.g. 'TextEdit', 'Notes', 'Word'"},
                        "text": {"type": "string", "description": "Exact text to type for open_app_and_type"},
                    },
                },
                "agent_id": {"type": "string", "description": "Agent ID (from broker_computer_use_status). If omitted, sends to first available agent."},
                "timeout": {"type": "integer", "description": "Max seconds to wait for result (default 120)"},
            },
            "required": ["connector", "action"],
        },
    },
    # ── Google Drive tools ──────────────────────────────────────────────────
    {
        "name": "broker_drive_upload",
        "description": "Upload a generated file (offer PDF, report XLSX/DOCX) to the broker Google Drive folder. Returns a shareable link. Use after broker_create_offer or broker_asf_summary to save the result to Drive.",
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "Local filename from output/ directory (e.g. Offer_CLI001_2026-03-10.pdf)"},
                "drive_filename": {"type": "string", "description": "Optional: name to use in Google Drive (default: same as local)"}
            },
            "required": ["filename"]
        }
    },
    {
        "name": "broker_drive_list",
        "description": "List files in the broker Google Drive folder. Returns names, sizes, modification dates, and shareable links.",
        "input_schema": {
            "type": "object",
            "properties": {
                "limit": {"type": "integer", "description": "Max files to return (default 20)"},
                "name_filter": {"type": "string", "description": "Optional filter by partial filename (e.g. Offer, ASF)"}
            }
        }
    },
    {
        "name": "broker_drive_get_link",
        "description": "Get a shareable Google Drive link for an already-uploaded file by its exact filename.",
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "Exact filename in Google Drive"}
            },
            "required": ["filename"]
        }
    },
    # ── SharePoint tools ─────────────────────────────────────────────────────
    {
        "name": "broker_sharepoint_upload",
        "description": "Upload a generated file to the broker SharePoint folder via Microsoft Graph API. Returns a SharePoint link. Use for companies on Microsoft 365.",
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "Local filename from output/ directory"},
                "sp_filename": {"type": "string", "description": "Optional: name to use in SharePoint"}
            },
            "required": ["filename"]
        }
    },
    {
        "name": "broker_sharepoint_list",
        "description": "List files in the broker SharePoint folder. Returns names, sizes, modification dates, and SharePoint links.",
        "input_schema": {
            "type": "object",
            "properties": {
                "limit": {"type": "integer", "description": "Max files to return (default 20)"},
                "name_filter": {"type": "string", "description": "Optional filter by partial filename"}
            }
        }
    },
    {
        "name": "broker_sharepoint_get_link",
        "description": "Get a SharePoint link for an already-uploaded file by its exact filename.",
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {"type": "string", "description": "Exact filename in SharePoint"}
            },
            "required": ["filename"]
        },
    },
    # ── RAG / Knowledge Base tools ─────────────────────────────────────────────
    {
        "name": "broker_search_knowledge",
        "description": (
            "Semantic search in the broker knowledge base (ChromaDB RAG). "
            "Use when broker asks vague/natural-language questions about: "
            "- What does a product cover? What are the exclusions? "
            "- Which insurer handles CMR claims? "
            "- What documents are needed for a claim at Allianz? "
            "- What ASF/BaFin class is this product? "
            "Returns top matching chunks with relevance scores. "
            "Categories: 'product', 'claims_guidance', 'compliance', 'faq_doc'."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Natural language question or keyword"},
                "category": {
                    "type": "string",
                    "description": "Optional filter: 'product', 'claims_guidance', 'compliance', 'faq_doc'",
                },
                "top_k": {"type": "integer", "description": "Number of results (default 5, max 10)"},
            },
            "required": ["query"],
        },
    },
    {
        "name": "broker_upload_document",
        "description": (
            "Upload a PDF or image to Anthropic Files API for persistent document analysis. "
            "Returns a file_id that can be reused in broker_analyze_document without re-uploading. "
            "Use for: scanned policies, damage photos, invoices, constatare amiabilă. "
            "Supported formats: PDF, PNG, JPG, WEBP, GIF."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path": {"type": "string", "description": "Absolute path to the file"},
                "doc_type": {
                    "type": "string",
                    "description": "Document type: 'policy', 'claim_photo', 'invoice', 'constatare', 'id_card', 'other'",
                },
                "description": {"type": "string", "description": "Short description for audit trail"},
            },
            "required": ["file_path"],
        },
    },
    {
        "name": "broker_analyze_document",
        "description": (
            "[BETA — în testare, verificați rezultatele manual] "
            "Analyze a document using Claude Vision. Accepts local file path OR file_id from broker_upload_document. "
            "Auto-extracts by doc_type: "
            "'policy' → policy number, dates, premium, coverage, exclusions; "
            "'invoice' → vendor, items, totals; "
            "'constatare' → both drivers, damage, signatures (handles handwriting in RO/DE/EN); "
            "'id_card' → name, CNP/ID, address. "
            "NOTE: doc_type='claim_photo' is DISABLED in production — under development for Phase 2. "
            "Returns structured markdown extraction — always verify extracted data before use."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "file_path_or_file_id": {
                    "type": "string",
                    "description": "Local path (e.g. /path/to/policy.pdf) OR file_id starting with 'file_'",
                },
                "question": {
                    "type": "string",
                    "description": "Specific question (optional — auto-prompt used if empty)",
                },
                "doc_type": {
                    "type": "string",
                    "description": "'policy', 'claim_photo', 'invoice', 'constatare', 'id_card', 'other'",
                },
            },
            "required": ["file_path_or_file_id"],
        },
    },
    {
        "name": "broker_kb_status",
        "description": "Show knowledge base status: total indexed chunks, breakdown by category. Admin/debug use.",
        "input_schema": {"type": "object", "properties": {}},
    },
    {
        "name": "broker_kb_reindex",
        "description": "Force re-index all knowledge sources (products, compliance maps, docs). Use after adding new products to DB or updating docs/ files.",
        "input_schema": {"type": "object", "properties": {}},
    },
    {
        "name": "broker_list_output_files",
        "description": "List all generated offer/report files (PDF, TXT, XLSX, DOCX) saved in the output directory. Shows filename, size, date. Use when broker asks about generated files or wants to clean up old ones.",
        "input_schema": {
            "type": "object",
            "properties": {
                "sort_by": {"type": "string", "description": "Sort by: 'date' (newest first, default) or 'name'"},
                "filter_ext": {"type": "string", "description": "Filter by extension: 'pdf', 'txt', 'xlsx', 'docx' or 'all' (default)"},
            },
        },
    },
    {
        "name": "broker_send_claim_questionnaire",
        "description": "Send a damage report questionnaire to a client for a specific claim. Auto-selects KFZ-Schaden or Haftpflicht template based on policy type. Sends via email + WhatsApp.",
        "input_schema": {
            "type": "object",
            "properties": {
                "claim_id": {"type": "string", "description": "Claim ID (e.g. CLM001)"},
                "template_id": {"type": "string", "description": "Optional: override template (tpl-kfz-schaden or tpl-haftpflicht)"},
            },
            "required": ["claim_id"],
        },
    },
    {
        "name": "broker_auto_send_questionnaires",
        "description": "Auto-detect ALL open claims without questionnaires and send the appropriate damage report form (KFZ-Schaden or Haftpflicht) to each client. Returns count of forms sent.",
        "input_schema": {"type": "object", "properties": {}},
    },
    {
        "name": "broker_check_form_status",
        "description": "Get summary of all form/questionnaire submissions: counts by status (sent, in_progress, completed), overdue forms, and recent submissions with client details.",
        "input_schema": {"type": "object", "properties": {}},
    },
    {
        "name": "broker_run_form_reminders",
        "description": "Run form follow-up: check all incomplete form submissions and send reminders to clients who haven't completed after 4, 8, or 12 days. Max 3 reminders per submission.",
        "input_schema": {"type": "object", "properties": {}},
    },
    {
        "name": "broker_form_daily_report",
        "description": "Generate daily form completion report: total forms, completion rate %, overdue count, recent submissions. Sends HTML report email to operator. Returns stats summary.",
        "input_schema": {"type": "object", "properties": {}},
    },
    {
        "name": "broker_oracle_dashboard",
        "description": "Connect to Oracle Database and return the full insurance broker dashboard: connection status, Oracle version, total clients, active policies, monthly premium revenue, open claims with details, form completion rate, and all database tables with row counts.",
        "input_schema": {"type": "object", "properties": {}},
    },
    {
        "name": "broker_oracle_query",
        "description": "Execute a read-only SQL query on the Oracle Database. Use standard Oracle SQL syntax. Returns columns and rows as JSON. Use for custom reports, analytics, filtering clients, policies, claims, or form submissions.",
        "input_schema": {
            "type": "object",
            "properties": {
                "sql": {"type": "string", "description": "Oracle SQL SELECT query to execute"},
            },
            "required": ["sql"],
        },
        "cache_control": {"type": "ephemeral"},
    },
    {
        "name": "broker_execute_command",
        "description": "Execute a CMD quick command. Available: CMD1 [ref] = check status, CMD2 [ref] = send reminder, CMD3 [name/email] = client history, CMD4 [ref] = resend form, CMD5 [ref] = close case, CMD6 [ref] \"question\" = send extra questions to client, CMD7 = Oracle sync, CMD8 = daily report. Ref = AKT-xxxx or FRM-xxxx.",
        "input_schema": {
            "type": "object",
            "properties": {
                "command": {"type": "string", "description": "Command code: CMD1, CMD2, CMD3, CMD4, CMD5, CMD6, CMD7"},
                "ref": {"type": "string", "description": "Reference: AKT-xxxx, FRM-xxxx, or client name/email for CMD3"},
            },
            "required": ["command"],
        },
    },
    {
        "name": "broker_list_forms",
        "description": "List all form submissions with status, client name, template, completeness. Use to see all questionnaires sent to clients.",
        "input_schema": {"type": "object", "properties": {}},
    },
    {
        "name": "broker_send_form_link",
        "description": "Send a new questionnaire link to a client. Creates a pre-submission and sends via email. Use when a broker wants to manually send a form to a specific client.",
        "input_schema": {
            "type": "object",
            "properties": {
                "template_id": {"type": "string", "description": "Template: tpl-kfz-schaden, tpl-haftpflicht, or tpl-maschinenbruch"},
                "client_name": {"type": "string", "description": "Client name"},
                "client_email": {"type": "string", "description": "Client email address"},
                "client_phone": {"type": "string", "description": "Client phone (optional, for WhatsApp)"},
            },
            "required": ["template_id", "client_name", "client_email"],
        },
    },
]

# ── Tool executor ─────────────────────────────────────────────────────────────
TOOL_DISPATCH = {
    "broker_search_clients":     search_clients_fn,
    "broker_get_client":         get_client_fn,
    "broker_create_client":      create_client_fn,
    "broker_update_client":      update_client_fn,
    "broker_delete_client":      delete_client_fn,
    "broker_search_products":    search_products_fn,
    "broker_compare_products":   compare_products_fn,
    "broker_create_offer":       create_offer_fn,
    "broker_list_offers":        list_offers_fn,
    "broker_send_offer_email":   send_offer_email_fn,
    "broker_get_renewals_due":   get_renewals_due_fn,
    "broker_list_policies":      list_policies_fn,
    "broker_log_claim":          log_claim_fn,
    "broker_get_claim_status":   get_claim_status_fn,
    "broker_asf_summary":        asf_summary_fn,
    "broker_bafin_summary":      bafin_summary_fn,
    "broker_check_rca_validity": check_rca_validity_fn,
    "broker_cross_sell":         cross_sell_fn,
    "broker_compare_premiums_live":   compare_premiums_live_fn,
    # Vehicle management
    "broker_add_vehicle":        add_vehicle_fn,
    "broker_list_vehicles":      list_vehicles_fn,
    "broker_search_vehicle":     search_vehicle_fn,
    "broker_get_vehicle":        get_vehicle_fn,
    "broker_update_vehicle":     update_vehicle_fn,
    "broker_delete_vehicle":     delete_vehicle_fn,
    "broker_compliance_check":   compliance_check_fn,
    "broker_save_conversation":  None,  # special — handled in agentic loop (needs session context)
    # Web automation (Playwright on Cloud Run — sync wrappers, run in thread pool)
    "broker_check_rca":          _playwright_check_rca_fn,
    "broker_browse_web":         _playwright_browse_web_fn,
    "broker_scrape_rca_prices":  _scrape_rca_prices_fn,
    # Computer use tools — status is sync, run_task is async (handled in agentic loop)
    "broker_computer_use_status": None,  # set after function definition below
    "broker_run_task":            None,  # async — handled in agentic loop
    # Google Drive tools
    "broker_drive_upload":        upload_to_drive_fn,
    "broker_drive_list":          list_drive_files_fn,
    "broker_drive_get_link":      get_drive_link_fn,
    # SharePoint tools
    "broker_sharepoint_upload":   sp_upload_fn,
    "broker_sharepoint_list":     sp_list_fn,
    "broker_sharepoint_get_link": sp_get_link_fn,
    # RAG / Knowledge base tools
    "broker_search_knowledge":    broker_search_knowledge_fn,
    "broker_upload_document":     broker_upload_document_fn,
    "broker_analyze_document":    broker_analyze_document_fn,
    "broker_kb_status":           broker_kb_status_fn,
    "broker_kb_reindex":          broker_kb_reindex_fn,
    # Output file management
    "broker_list_output_files":   None,  # handled inline — uses OUTPUT_DIR
    # Questionnaire / Form management tools
    "broker_send_claim_questionnaire": None,  # async — calls API endpoint
    "broker_auto_send_questionnaires": None,  # async — calls API endpoint
    "broker_check_form_status":        None,  # async — calls API endpoint
    "broker_run_form_reminders":       None,  # async — calls function directly
    "broker_form_daily_report":        None,  # async — calls function directly
    "broker_oracle_dashboard":         None,  # async — calls Oracle API
    "broker_oracle_query":             None,  # async — calls Oracle API
    "broker_execute_command":          None,  # async — calls CMD API
    "broker_list_forms":               None,  # async — calls forms API
    "broker_send_form_link":           None,  # async — calls send-link API
}

def execute_tool(tool_name: str, tool_input: dict) -> str:
    fn = TOOL_DISPATCH.get(tool_name)
    if fn is None and tool_name in TOOL_DISPATCH:
        # Async tool — should be handled by execute_tool_async
        return f"[async tool '{tool_name}' — use execute_tool_async]"
    if not fn:
        return (f"Tool '{tool_name}' does not exist. "
                f"Available tools: {', '.join(TOOL_DISPATCH.keys())}")
    try:
        result = fn(**tool_input)
        return result if result else "Operation completed successfully."
    except TypeError as e:
        # Missing or wrong arguments — give helpful hint
        import inspect
        sig = inspect.signature(fn)
        return (f"Incorrect parameters for {tool_name}. "
                f"Required parameters: {list(sig.parameters.keys())}. "
                f"Received: {list(tool_input.keys())}. Error: {e}")
    except Exception as e:
        err = str(e)
        # Return friendly message with context so Alex can recover
        return (f"An issue occurred with {tool_name}: {err[:200]}. "
                f"Input used: {tool_input}. "
                f"Suggestion: check if the client/product ID is correct.")


# ── Computer Use tool implementations (async) ─────────────────────────────────

def _cu_computer_use_status_fn(**kwargs) -> str:
    """Returns status of connected local agents.

    Reads directly from cu_state._cu_agents (shared module, same process).
    No HTTP call needed — app.py and main.py share the same cu_state module
    because they run in the same uvicorn process via mount_chainlit.
    """
    from datetime import datetime as _dt2
    online = []
    for agent_id, info in _cu_agents.items():
        try:
            last = _dt2.fromisoformat(info["last_seen"])
            delta = (_dt2.utcnow() - last).total_seconds()
            if delta < 120:
                online.append({**info, "online": True, "seconds_ago": round(delta)})
        except Exception:
            pass

    if not online:
        return (
            "⚠️ **No local agent connected.**\n\n"
            "To use computer use, start the local agent on the employee's machine:\n"
            "```\n"
            "cd alex-local-agent\n"
            "python main.py start\n"
            "```"
        )

    lines = ["## 🤖 Local Agent — Status\n"]
    for info in online:
        secs = info.get("seconds_ago", "?")
        lines.append(f"**Agent:** `{info.get('agent_id', '?')}`")
        lines.append(f"- Status: 🟢 Online (last seen {secs}s ago)")
        lines.append(f"- Platform: {info.get('platform', 'N/A')}")
        connectors = info.get("connectors", [])
        lines.append(f"- Available connectors: {', '.join(connectors) if connectors else 'none'}")
        lines.append("")

    return "\n".join(lines)


async def _cu_run_task_async(
    connector: str,
    action: str,
    params: dict = None,
    agent_id: str = None,
    timeout: int = 120,
    credentials: dict = None,
) -> str:
    """Async implementation of broker_run_task — enqueues task and waits for result."""
    params = params or {}
    credentials = credentials or {}

    # Determine target agent — read directly from cu_state (same process, no HTTP needed)
    if not agent_id:
        from datetime import datetime as _dt3
        for info in _cu_agents.values():
            try:
                last = _dt3.fromisoformat(info["last_seen"])
                delta = (_dt3.utcnow() - last).total_seconds()
                if delta < 120:
                    agent_id = info["agent_id"]
                    break
            except Exception:
                pass

    if not agent_id:
        return (
            "⚠️ **No local agent online.**\n\n"
            "Start the local agent with:\n"
            "```\ncd /Users/andreibotescu/Desktop/insurance-broker-agent/alex-local-agent\n"
            "python main.py start\n```"
        )

    # Enqueue task
    task_id = _cu_enqueue_task(
        connector=connector,
        action=action,
        params=params,
        agent_id=agent_id,
        timeout=timeout,
        credentials=credentials,
    )

    # Notify user we're waiting
    connector_emoji = {"cedam": "🚗", "web_generic": "🌐", "desktop_generic": "🖥️",
                       "anthropic_computer_use": "🤖"}.get(connector, "⚙️")
    action_desc = {
        "check_rca": f"RCA check for {params.get('plate', '')}",
        "extract": f"extracting: {params.get('query', '')[:60]}",
        "navigate": f"navigating to {params.get('url', '')}",
        "fill_form": "filling form",
        "screenshot": "screen capture",
        "read_screen": "reading screen",
        "login": "authentication",
    }.get(action, action)

    wait_msg = f"{connector_emoji} **Running task on local agent** ({connector} / {action_desc})..."

    # Wait for result with timeout
    result = await _cu_wait_result(task_id, timeout=timeout + 10)

    if result.get("captcha_detected") and not result.get("success"):
        visible_used = result.get("visible_browser_used", False)
        if visible_used:
            return (
                "🔒 **CAPTCHA not solved**\n\n"
                "The browser opened on your computer. "
                "Complete the CAPTCHA in the browser window, "
                "then tell me to retry the check."
            )
        else:
            return (
                "🔒 **AIDA portal requires CAPTCHA**\n\n"
                "The local agent will open a visible browser on your computer. "
                "Check if the agent is running with:\n"
                "```\ncd /Users/andreibotescu/Desktop/insurance-broker-agent/alex-local-agent\npython main.py start\n```\n"
                "If it's running, try again — the browser will open and you can complete the CAPTCHA manually."
            )

    if not result.get("success"):
        error = result.get("error", "Unknown error")
        return f"❌ **Task failed** ({connector}/{action})\n\nError: {error}"

    # Format result nicely
    if action == "check_rca":
        return _format_rca_result(result, params.get("plate", ""))
    elif action == "screenshot":
        # Return info about screenshot (base64 would be too long for chat)
        size = result.get("size_bytes", 0)
        return f"📸 **Screenshot captured** ({size // 1024} KB)\n\nI can analyze the screen if you want — tell me what to look for."
    elif action == "extract":
        data = result.get("data") or result.get("result", {})
        if isinstance(data, str):
            return f"📄 **Extracted data:**\n\n{data[:2000]}"
        elif isinstance(data, (list, dict)):
            return f"📄 **Extracted data:**\n\n```json\n{json.dumps(data, indent=2, ensure_ascii=False)[:2000]}\n```"
        return f"📄 **Result:** {str(result)[:1000]}"
    else:
        # Generic success response
        data = result.get("data") or result.get("result") or result
        if isinstance(data, str):
            return f"✅ **Task completed** ({connector}/{action})\n\n{data[:1500]}"
        return f"✅ **Task completed** ({connector}/{action})\n\n```\n{json.dumps(data, indent=2, ensure_ascii=False, default=str)[:1500]}\n```"


def _format_rca_result(result: dict, plate: str) -> str:
    """Format RCA check result as readable markdown."""
    if not result.get("data_found"):
        return (
            f"❌ **No active RCA** for plate number **{plate}**\n\n"
            f"The client needs to purchase an RCA policy."
        )

    valid = result.get("rca_valid")
    expiry = result.get("expiry_date", "N/A")
    days = result.get("days_until_expiry")
    insurer = result.get("insurer", "N/A")
    policy_no = result.get("policy_number", "N/A")

    status_icon = "✅" if valid else "❌"
    status_text = "**VALID**" if valid else "**EXPIRED**"

    urgency = ""
    if days is not None and days <= 30:
        urgency = f"\n\n⚠️ **WARNING: Policy expires in {days} days!** Contact the client for renewal."
    elif days is not None and days <= 60:
        urgency = f"\n\n📅 Policy expires in {days} days — prepare renewal."

    return (
        f"## {status_icon} RCA Check — {plate}\n\n"
        f"- **Status:** {status_text}\n"
        f"- **Insurer:** {insurer}\n"
        f"- **Policy number:** {policy_no}\n"
        f"- **Expiry date:** {expiry}\n"
        f"- **Days remaining:** {days if days is not None else 'N/A'}"
        f"{urgency}"
    )


# ── Register async computer use tools in dispatch (after functions are defined) ─
TOOL_DISPATCH["broker_computer_use_status"] = _cu_computer_use_status_fn
# broker_run_task stays None — handled specially in agentic loop (await)


# ── Export helpers ────────────────────────────────────────────────────────────

def export_to_xlsx(content: str, title: str) -> Path:
    """Convert text content to XLSX with formatting."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        return None

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:31]

    # Header style
    header_fill = PatternFill(start_color="1a365d", end_color="1a365d", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)

    lines = content.split("\n")
    row = 1
    for line in lines:
        if line.startswith("|") and "|" in line[1:]:
            # Table row
            cells = [c.strip() for c in line.split("|")[1:-1]]
            for col, cell_text in enumerate(cells, 1):
                cell = ws.cell(row=row, column=col, value=cell_text)
                if row <= 2:
                    cell.fill = header_fill
                    cell.font = header_font
                cell.alignment = Alignment(wrap_text=True)
            ws.column_dimensions[chr(64 + min(col, 26))].width = 20
        elif line.startswith("#"):
            cell = ws.cell(row=row, column=1, value=line.lstrip("#").strip())
            cell.font = Font(bold=True, size=12)
        elif line.strip() and not line.startswith("---"):
            ws.cell(row=row, column=1, value=line)

        if line.strip() or line.startswith("|"):
            row += 1

    out_path = OUTPUT_DIR / f"{title.replace(' ', '_')}_{date.today().isoformat()}.xlsx"
    wb.save(str(out_path))
    return out_path


def export_to_docx(content: str, title: str) -> Path:
    """Convert text content to DOCX."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = Document()

    # Title
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    lines = content.split("\n")
    table_rows = []
    in_table = False

    for line in lines:
        if line.startswith("|") and "|" in line[1:]:
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if not line.replace("|", "").replace("-", "").replace(" ", ""):
                continue  # separator line
            table_rows.append(cells)
            in_table = True
        else:
            if in_table and table_rows:
                # Flush table
                if len(table_rows) > 1:
                    t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    t.style = "Table Grid"
                    for r_idx, row_data in enumerate(table_rows):
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx < len(t.rows[r_idx].cells):
                                t.rows[r_idx].cells[c_idx].text = cell_text
                                if r_idx == 0:
                                    t.rows[r_idx].cells[c_idx].paragraphs[0].runs[0].font.bold = True if t.rows[r_idx].cells[c_idx].paragraphs[0].runs else True
                doc.add_paragraph()
                table_rows = []
                in_table = False

            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line.strip("**"))
                run.bold = True
            elif line.strip() and not line.startswith("---"):
                doc.add_paragraph(line)

    out_path = OUTPUT_DIR / f"{title.replace(' ', '_')}_{date.today().isoformat()}.docx"
    doc.save(str(out_path))
    return out_path


def export_to_pdf(content: str, title: str) -> Path:
    """Convert text/markdown content to PDF via WeasyPrint."""
    try:
        from weasyprint import HTML
    except ImportError:
        return None

    # Convert markdown-like content to HTML
    lines = content.split("\n")
    html_lines = ["<html><head><style>",
                  "body { font-family: Arial, sans-serif; margin: 40px; color: #333; }",
                  "h1 { color: #1a365d; border-bottom: 2px solid #1a365d; }",
                  "h2 { color: #2d3748; }",
                  "table { border-collapse: collapse; width: 100%; margin: 15px 0; }",
                  "th { background: #1a365d; color: white; padding: 8px; text-align: left; }",
                  "td { padding: 8px; border: 1px solid #e2e8f0; }",
                  "tr:nth-child(even) { background: #f7fafc; }",
                  ".footer { color: #718096; font-size: 11px; margin-top: 30px; border-top: 1px solid #e2e8f0; padding-top: 10px; }",
                  "</style></head><body>"]

    in_table = False
    table_header_done = False

    for line in lines:
        if line.startswith("# "):
            html_lines.append(f"<h1>{line[2:]}</h1>")
        elif line.startswith("## "):
            html_lines.append(f"<h2>{line[3:]}</h2>")
        elif line.startswith("### "):
            html_lines.append(f"<h3>{line[4:]}</h3>")
        elif line.startswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if all(c.replace("-", "").replace(" ", "") == "" for c in cells):
                continue  # separator
            if not in_table:
                html_lines.append("<table>")
                in_table = True
                table_header_done = False
            if not table_header_done:
                html_lines.append("<tr>" + "".join(f"<th>{c}</th>" for c in cells) + "</tr>")
                table_header_done = True
            else:
                html_lines.append("<tr>" + "".join(f"<td>{c}</td>" for c in cells) + "</tr>")
        else:
            if in_table:
                html_lines.append("</table>")
                in_table = False
                table_header_done = False
            if line.startswith("---"):
                html_lines.append("<hr>")
            elif line.strip():
                # Bold text
                line = line.replace("**", "<strong>", 1).replace("**", "</strong>", 1) if "**" in line else line
                html_lines.append(f"<p>{line}</p>")

    if in_table:
        html_lines.append("</table>")

    html_lines.append(f'<div class="footer">Generated by Alex — Insurance Broker AI | {date.today().strftime("%d %B %Y")}</div>')
    html_lines.append("</body></html>")

    html_content = "\n".join(html_lines)
    out_path = OUTPUT_DIR / f"{title.replace(' ', '_')}_{date.today().isoformat()}.pdf"

    try:
        HTML(string=html_content).write_pdf(str(out_path))
        return out_path
    except Exception:
        return None


async def send_export_files(content: str, base_title: str):
    """Generate and send PDF, XLSX, DOCX as downloadable attachments."""
    elements = []
    generated = []

    xlsx_path = export_to_xlsx(content, base_title)
    if xlsx_path and xlsx_path.exists():
        elements.append(cl.File(name=xlsx_path.name, path=str(xlsx_path), display="side"))
        generated.append("XLSX")

    docx_path = export_to_docx(content, base_title)
    if docx_path and docx_path.exists():
        elements.append(cl.File(name=docx_path.name, path=str(docx_path), display="side"))
        generated.append("DOCX")

    pdf_path = export_to_pdf(content, base_title)
    if pdf_path and pdf_path.exists():
        elements.append(cl.File(name=pdf_path.name, path=str(pdf_path), display="side"))
        generated.append("PDF")

    if elements:
        await cl.Message(
            content=f"📎 **Export ready:** {' · '.join(generated)} — click to download",
            elements=elements,
            author="Alex 🤖"
        ).send()


# ── Dashboard helpers ─────────────────────────────────────────────────────────
def get_dashboard_stats() -> dict:
    try:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        today = date.today().isoformat()
        in_30 = (date.today() + timedelta(days=30)).isoformat()
        in_7  = (date.today() + timedelta(days=7)).isoformat()

        active      = conn.execute("SELECT COUNT(*) as n FROM policies WHERE status='active'").fetchone()["n"]
        expiring7   = conn.execute("SELECT COUNT(*) as n FROM policies WHERE status='active' AND end_date<=? AND end_date>=?", (in_7, today)).fetchone()["n"]
        expiring30  = conn.execute("SELECT COUNT(*) as n FROM policies WHERE status='active' AND end_date<=? AND end_date>=?", (in_30, today)).fetchone()["n"]
        clients     = conn.execute("SELECT COUNT(*) as n FROM clients").fetchone()["n"]
        open_claims = conn.execute("SELECT COUNT(*) as n FROM claims WHERE status='open'").fetchone()["n"]
        offers      = conn.execute("SELECT COUNT(*) as n FROM offers").fetchone()["n"]
        conn.close()
        return {"active_policies": active, "expiring_7": expiring7, "expiring_30": expiring30,
                "clients": clients, "open_claims": open_claims, "offers_sent": offers}
    except Exception:
        return {}


async def process_uploaded_file(file: cl.File) -> str:
    """Process an uploaded file with Claude Vision. Returns analysis text."""
    name_lower = file.name.lower()
    is_image = any(name_lower.endswith(ext) for ext in [".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp"])
    is_pdf = name_lower.endswith(".pdf")
    is_text = any(name_lower.endswith(ext) for ext in [".txt", ".md", ".csv"])

    if not (is_image or is_pdf or is_text):
        return f"File {file.name} received — format not supported for auto-analysis. Please describe its content."

    # Chainlit 2.x: uploaded files may have content=None, read from path instead
    file_bytes = file.content
    if not file_bytes and file.path:
        file_bytes = Path(file.path).read_bytes()
    if not file_bytes:
        return f"Could not read file content for {file.name}."

    if is_text:
        text_content = file_bytes.decode("utf-8", errors="replace")
        return f"[Text file: {file.name}]\n\n{text_content[:3000]}"

    if is_image:
        media_type = (
            "image/jpeg" if name_lower.endswith((".jpg", ".jpeg")) else
            "image/png" if name_lower.endswith(".png") else
            "image/webp" if name_lower.endswith(".webp") else
            "image/jpeg"
        )
    else:
        media_type = "application/pdf"

    analysis_prompt = """Analyze this document carefully.

IMPORTANT DISTINCTION:
- The **CLIENT/PATIENT/SUBJECT** is the person this document is ABOUT (the insured person, patient, vehicle owner, policy holder)
- The **ISSUER/CLINIC/INSURER** is the organization that ISSUED or SIGNED the document

Extract in this exact format:

**Document Type:** (ID card / medical referral / insurance policy / accident report / invoice / lab results / other)

**CLIENT DATA** (the person this document is about — NOT the issuing organization):
- Full Name: [name of the person, NOT the clinic/company/issuer]
- Phone: [client's personal phone, NOT clinic reception number]
- Email: [client's email if present]
- Address: [client's address]
- Date of Birth / CNP: [if present]
- ID/Policy Number: [if present]

**ISSUER DATA** (clinic, insurer, company that issued the document):
- Issuer Name: [clinic name, insurer name, company]
- Issuer Phone: [clinic/company contact]
- Issuer Address: [clinic/company address]

**OTHER DETAILS:**
- Dates: [relevant dates]
- Amounts: [any monetary amounts in RON/EUR]
- Vehicle details: [if applicable]
- Diagnosis/Coverage/Damage: [key content]

**Suggested Insurance Type:** [HEALTH / LIFE / RCA / CASCO / PAD / other based on document context]

Be precise. Never confuse issuer contact data with client personal data."""

    try:
        encoded = base64.standard_b64encode(file_bytes).decode("utf-8")
        # Claude supports images directly; PDFs via base64 document blocks
        if is_image:
            content_block = {
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": encoded,
                },
            }
        else:
            # PDF — send as document block (supported in claude-3+ models)
            content_block = {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": encoded,
                },
            }
        response = await _asyncio.get_event_loop().run_in_executor(
            None,
            lambda: client.messages.create(
                model=MODEL,
                max_tokens=1024,
                messages=[{
                    "role": "user",
                    "content": [
                        content_block,
                        {"type": "text", "text": analysis_prompt},
                    ],
                }],
            )
        )
        analysis = ""
        for block in response.content:
            if hasattr(block, "text"):
                analysis += block.text
        return f"[Document analyzed: {file.name}]\n\n{analysis}" if analysis else f"Could not extract content from {file.name}."
    except Exception as e:
        err_str = str(e)
        return (f"⚠️ Could not analyze {file.name}. "
                f"(Error: {err_str[:200]})")


# ── Auth callback (only active when CHAINLIT_AUTH_SECRET is set) ──────────────
if os.environ.get("CHAINLIT_AUTH_SECRET") and ADMIN_ENABLED:
    @cl.password_auth_callback
    def auth_callback(username: str, password: str):
        user = get_user_by_email(username.strip())
        if not user or not user["is_active"]:
            return None
        if not verify_password(password, user["hashed_password"]):
            return None
        return cl.User(
            identifier=user["email"],
            metadata={
                "user_id": user["id"],
                "role": user["role"],
                "company_id": user["company_id"],
                "full_name": user["full_name"] or user["email"],
            }
        )


# ── Helper: load tools + show dashboard welcome ───────────────────────────────

async def _init_session_tools():
    """Load allowed_tools and user_meta into session. Returns user metadata dict."""
    if ADMIN_ENABLED:
        app_user = cl.user_session.get("user")
        if app_user:
            meta = app_user.metadata or {}
            allowed = get_user_tools(meta.get("user_id"), meta.get("role"))
            cl.user_session.set("allowed_tools", allowed)
            cl.user_session.set("user_meta", meta)
            return meta
        else:
            # No login active — use demo user so history/conversations still work
            _demo_meta = {"user_id": "demo", "role": "admin", "full_name": "Broker"}
            cl.user_session.set("allowed_tools", None)  # all tools allowed
            cl.user_session.set("user_meta", _demo_meta)
            return _demo_meta
    return {}


async def _build_sidebar_content(user_id: str) -> str:
    """Build markdown content for the saved conversations sidebar panel."""
    try:
        clients = list_clients_with_conversations(user_id) if ADMIN_ENABLED else []
    except Exception:
        clients = []

    if not clients:
        return (
            "## 📁 Saved Conversations\n\n"
            "_No saved conversations yet._\n\n"
            "**How to save:**\n"
            "1. Chat with Alex about a client\n"
            "2. Say **\"save conversation\"**\n"
            "3. It appears here grouped by client\n\n"
            "_Or after 3+ messages the button appears automatically._"
        )

    lines = ["## 📁 Saved Conversations\n"]
    for c in clients[:30]:
        cname = c["client_name"] if c["client_name"] != "__unlinked__" else "No client linked"
        conv_count = c.get("conv_count", 0)
        lines.append(f"### 👤 {cname} ({conv_count})")
        try:
            convs = list_conversations_for_client(user_id, c["client_id"])
        except Exception:
            convs = []
        for conv in convs[:8]:
            msgs  = conv.get("message_count", 0)
            upd   = conv.get("updated_at", "")[:10]
            label = conv.get("title", "Conversation")
            detail = f"{msgs} msg · {upd}" if msgs else upd
            lines.append(f"- **{label}**  \n  _{detail}_")
        lines.append("")
    lines.append("---\n_Say 'save conversation' to add one._")
    return "\n".join(lines)


async def _refresh_sidebar(user_id: str) -> None:
    """Deprecated stub — sidebar is now attached to the welcome message element."""
    pass  # See _show_dashboard_welcome which sends the sidebar Text element inline


async def _show_dashboard_welcome(user_id: str | None = None, full_name: str | None = None):
    """Send the portfolio dashboard welcome message with optional history shortcut."""
    stats = get_dashboard_stats()
    alerts = ""
    if stats.get("expiring_7", 0) > 0:
        n = stats['expiring_7']
        alerts = f"\n> ⚠️ **{n} {'policy' if n == 1 else 'policies'} expiring within 7 days!** Try: *'Show urgent renewals'*"

    # Pending approvals count
    pending_approvals = 0
    try:
        import sqlite3 as _sq3
        _adb = _sq3.connect(str(Path(__file__).parent / "mcp-server" / "insurance_broker.db"))
        pending_approvals = _adb.execute("SELECT COUNT(*) FROM approval_queue WHERE status = 'pending'").fetchone()[0]
        _adb.close()
    except Exception:
        pass
    approval_line = ""
    if pending_approvals > 0:
        approval_line = f"\n| 📬 Pending Approval | **{pending_approvals}** | [🔗 Open Dashboard](/dashboard/approvals) |"

    greeting_line = f"👋 Hello, **{full_name}**!\n\n" if full_name else ""

    welcome = f"""{greeting_line}## 📊 Portfolio Dashboard — {date.today().strftime('%d %B %Y')}

| Metric | Value | Action |
|---|---|---|
| 🟢 Active Policies | **{stats.get('active_policies', 0)}** | [View →](/dashboard/database) |
| ⚠️ Expiring in 7 Days | **{stats.get('expiring_7', 0)}** | [Renewals →](/dashboard/approvals) |
| 📅 Expiring in 30 Days | **{stats.get('expiring_30', 0)}** | [Policies →](/dashboard/database) |
| 👥 Clients | **{stats.get('clients', 0)}** | [Clients →](/dashboard/database) |
| 📋 Open Claims | **{stats.get('open_claims', 0)}** | [Claims →](/dashboard/database) |
| 📄 Offers Generated | **{stats.get('offers_sent', 0)}** | [Offers →](/dashboard/database) |{approval_line}
{alerts}

How can I help you? *(upload a document, ask about clients, renewals, compliance...)*"""

    actions = []
    # Dashboard button — always visible
    actions.append(cl.Action(
        name="open_dashboard",
        label=f"📬 Dashboard ({pending_approvals})" if pending_approvals > 0 else "📬 Dashboard",
        payload={"url": "/dashboard/approvals"},
    ))
    # Quick commands menu
    actions.append(cl.Action(
        name="quick_commands",
        label="📋 Quick Commands",
        payload={},
    ))
    # Public forms link
    actions.append(cl.Action(
        name="open_dashboard",
        label="📝 Insurance Forms",
        payload={"url": "/forms"},
    ))
    # Agent local status
    from datetime import datetime as _dt_check
    agent_online = False
    for _ag_id, _ag_info in _cu_agents.items():
        try:
            _last = _dt_check.fromisoformat(_ag_info["last_seen"])
            if (_dt_check.utcnow() - _last).total_seconds() < 120:
                agent_online = True
                break
        except Exception:
            pass
    actions.append(cl.Action(
        name="agent_local_toggle",
        label="🟢 Local Agent ON" if agent_online else "🔴 Local Agent OFF",
        payload={"online": agent_online},
    ))
    # History buttons — show even without login (use fallback user_id)
    _hist_user_id = user_id or "demo"
    if ADMIN_ENABLED:
        actions.append(cl.Action(
            name="open_history",
            label="📁 Saved Conversations",
            payload={"user_id": _hist_user_id},
        ))
        actions.append(cl.Action(
            name="search_conversations",
            label="🔍 Search Conversations",
            payload={},
        ))

    # Attach sidebar as a named Text element — display="side" opens the right panel
    # and stays open as long as the message is visible (persists in Chainlit 2.x)
    elements = []
    if ADMIN_ENABLED:
        sidebar_content = await _build_sidebar_content(_hist_user_id)
        elements.append(cl.Text(
            name="📁 Saved Conversations",
            content=sidebar_content,
            display="side",
        ))

    await cl.Message(content=welcome, actions=actions, elements=elements, author="Alex 🤖").send()


# ── Conversation picker helpers (client-based) ────────────────────────────────

async def _show_client_history_picker(user_id: str):
    """Show clients that have saved conversations — user picks a client to browse."""
    clients = list_clients_with_conversations(user_id) if ADMIN_ENABLED else []

    if not clients:
        await cl.Message(
            content="No saved conversations yet. Start chatting — Alex will save conversations automatically when you link them to a client.",
            author="Alex 🤖",
        ).send()
        return

    actions = []
    for c in clients[:20]:
        conv_word = "conversation" if c["conv_count"] == 1 else "conversations"
        actions.append(cl.Action(
            name="select_client_history",
            label=f"👤 {c['client_name']} ({c['conv_count']} {conv_word})",
            payload={"client_id": c["client_id"], "client_name": c["client_name"]},
        ))

    await cl.Message(
        content="**Conversation history** — select a client to see past conversations:",
        actions=actions,
        author="Alex 🤖",
    ).send()


async def _show_conversation_picker_for_client(user_id: str, client_id: str, client_name: str):
    """Show conversations for a specific client."""
    conversations = list_conversations_for_client(user_id, client_id) if ADMIN_ENABLED else []

    actions = []
    for c in conversations[:20]:
        msgs = c.get("message_count", 0)
        msg_word = "message" if msgs == 1 else "messages"
        updated = c.get("updated_at", "")[:10]  # just the date part
        label = f"🗨️ {c['title']}"
        if msgs:
            label += f"  ({msgs} {msg_word}, {updated})"
        actions.append(cl.Action(
            name="resume_conversation",
            label=label,
            payload={"conversation_id": c["id"], "title": c["title"]},
        ))

    actions.append(cl.Action(
        name="new_conversation_for_client",
        label=f"➕ New conversation about {client_name}",
        payload={"client_id": client_id, "client_name": client_name},
    ))

    name_display = client_name if client_id != "__unlinked__" else "unlinked conversations"
    await cl.Message(
        content=f"**{name_display}** — pick a conversation or start a new one:",
        actions=actions,
        author="Alex 🤖",
    ).send()


# ── Legacy project picker (still works for backwards compat) ──────────────────

async def _show_project_picker(user_id: str, full_name: str):
    """Redirect to client-based history picker."""
    await _show_client_history_picker(user_id)


async def _show_conversation_picker(user_id: str, project_id: int, project_name: str):
    """Legacy shim — kept for backwards compat."""
    conversations = list_conversations(user_id, project_id) if ADMIN_ENABLED else []
    actions = []
    for c in conversations[:20]:
        actions.append(cl.Action(
            name="resume_conversation",
            label=f"🗨️ {c['title']}",
            payload={"conversation_id": c["id"], "title": c["title"]},
        ))
    actions.append(cl.Action(
        name="new_conversation",
        label="➕ New conversation",
        payload={"project_id": project_id, "project_name": project_name},
    ))
    await cl.Message(
        content=f"**{project_name}** — pick a conversation or start a new one:",
        actions=actions,
        author="Alex 🤖",
    ).send()


async def _render_history_to_ui(history: list[dict]):
    """Re-display stored messages in the UI. Only text blocks shown (no raw tool JSON)."""
    for msg in history:
        role    = msg.get("role", "user")
        content = msg.get("content", "")

        if isinstance(content, list):
            # Multi-block assistant message — show only text parts
            text = "\n".join(
                block["text"]
                for block in content
                if isinstance(block, dict) and block.get("type") == "text" and block.get("text")
            )
        else:
            text = str(content)

        if not text.strip():
            continue

        author = "You" if role == "user" else "Alex 🤖"
        await cl.Message(content=text, author=author).send()


# ── Starters (quick-action chips in the input area) ──────────────────────────
@cl.set_starters
async def set_starters():
    return [
        cl.Starter(
            label="⚠️ Urgent Renewals",
            message="Show policies expiring in 30 days and generate renewal emails",
        ),
        cl.Starter(
            label="📊 Full Report",
            message="Generate the full report: renewals + cross-sell + compliance + claims",
        ),
        cl.Starter(
            label="🔍 Check RCA",
            message="Check RCA for plate number: ",
        ),
        cl.Starter(
            label="👥 Search Clients",
            message="Search all active clients",
        ),
    ]


# ── Chainlit lifecycle ────────────────────────────────────────────────────────
@cl.on_chat_start
async def on_chat_start():
    # Initialise session state
    cl.user_session.set(_SK_HISTORY,    [])
    cl.user_session.set(_SK_CONV_ID,    None)
    cl.user_session.set(_SK_PROJECT_ID, None)
    cl.user_session.set(_SK_CLIENT_ID,  None)
    cl.user_session.set(_SK_TITLE_SET,  False)
    cl.user_session.set(_SK_SAVE_NUDGE, False)

    meta = await _init_session_tools()
    user_id   = meta.get("user_id")
    full_name = meta.get("full_name") or "Broker"

    # Send welcome ONLY ONCE per thread.
    # on_chat_start fires on every reconnect/refresh — don't duplicate.
    _already = cl.user_session.get("_welcome_sent")
    if not _already:
        cl.user_session.set("_welcome_sent", True)
        await _show_dashboard_welcome(user_id=user_id, full_name=full_name)
        # Open sidebar with saved conversations (non-blocking, cosmetic)
        if user_id:
            await _refresh_sidebar(user_id)


# ── Action callbacks ──────────────────────────────────────────────────────────

@cl.action_callback("select_project")
async def on_select_project(action: cl.Action):
    await action.remove()
    meta       = cl.user_session.get("user_meta", {})
    user_id    = meta.get("user_id")
    project_id   = action.payload["project_id"]
    project_name = action.payload["project_name"]

    cl.user_session.set(_SK_PROJECT_ID, project_id)
    await _show_conversation_picker(user_id, project_id, project_name)


@cl.action_callback("new_project")
async def on_new_project(action: cl.Action):
    await action.remove()
    meta       = cl.user_session.get("user_meta", {})
    user_id    = meta.get("user_id")
    company_id = meta.get("company_id")

    res = await cl.AskUserMessage(
        content="Enter a name for your new project:",
        timeout=120,
    ).send()
    if not res:
        await cl.Message(content="Project creation cancelled.", author="Alex 🤖").send()
        return

    name = res.get("output", "").strip()[:100]
    if not name:
        await cl.Message(content="Project name cannot be empty.", author="Alex 🤖").send()
        return

    try:
        project = create_project(user_id, company_id, name)
    except Exception:
        await cl.Message(
            content=f'A project named **"{name}"** already exists. Please choose a different name.',
            author="Alex 🤖",
        ).send()
        return

    cl.user_session.set(_SK_PROJECT_ID, project["id"])
    await _show_conversation_picker(user_id, project["id"], name)


@cl.action_callback("no_project")
async def on_no_project(action: cl.Action):
    """Legacy — kept for backwards compat with any in-flight sessions."""
    await action.remove()
    cl.user_session.set(_SK_PROJECT_ID, None)
    cl.user_session.set(_SK_CONV_ID, None)
    meta = cl.user_session.get("user_meta", {})
    await _show_dashboard_welcome(
        user_id=meta.get("user_id"),
        full_name=meta.get("full_name"),
    )


@cl.action_callback("open_dashboard")
async def on_open_dashboard(action: cl.Action):
    """Open a dashboard or page link in a new tab."""
    url = action.payload.get("url", "/dashboard/approvals")
    if url == "/forms":
        await cl.Message(
            content=(
                "📝 **Versicherungsformulare**\n\n"
                "👉 [**Formulare öffnen →**](/forms)\n\n"
                "Verfügbare Fragebögen:\n"
                "- 🚗 [KFZ-Schadenmeldung](/forms/tpl-kfz-schaden) — Fahrzeugschaden melden\n"
                "- 📋 [Haftpflicht-Schadenanzeige](/forms/tpl-haftpflicht) — Haftpflichtschaden melden\n\n"
                "📊 [Dashboard →](/dashboard/forms) | ❓ [FAQ →](/forms/faq)"
            ),
            author="Alex 🤖",
        ).send()
    else:
        await cl.Message(
            content=(
                "📬 **Approvals Dashboard**\n\n"
                "Open the dashboard in a new tab:\n\n"
                f"👉 [**Open Approvals Dashboard →**]({url})\n\n"
                "There you can:\n"
                "- ✅ Approve and send emails to clients\n"
                "- 📱 Send WhatsApp messages\n"
                "- ✏️ Edit emails before sending\n"
                "- ❌ Reject unsuitable proposals"
            ),
            author="Alex 🤖",
        ).send()


@cl.action_callback("quick_commands")
async def on_quick_commands(action: cl.Action):
    """Show a comprehensive list of available commands."""
    all_cmds = [
        cl.Action(name="cmd_daily_brief", label="☀️ Daily Briefing", payload={}),
        cl.Action(name="cmd_renewals", label="⚠️ Urgent Renewals", payload={}),
        cl.Action(name="cmd_claims", label="🔧 Claims Follow-up", payload={}),
        cl.Action(name="cmd_cross_sell", label="💰 Cross-sell", payload={}),
        cl.Action(name="cmd_search_client", label="👥 Search Client", payload={}),
        cl.Action(name="cmd_check_rca", label="🔍 Check RCA", payload={}),
        cl.Action(name="cmd_send_questionnaire", label="📋 Send Questionnaire", payload={}),
        cl.Action(name="cmd_check_forms", label="📝 Form Status", payload={}),
        cl.Action(name="cmd_form_report", label="📊 Form Report", payload={}),
        cl.Action(name="cmd_form_reminders", label="🔔 Form Reminders", payload={}),
        cl.Action(name="cmd_generate_offers", label="📄 Generate Offers", payload={}),
        cl.Action(name="cmd_send_emails", label="✉️ Send Emails", payload={}),
        cl.Action(name="cmd_oracle_status", label="🗄️ Oracle DB", payload={}),
        cl.Action(name="cmd_knowledge_status", label="🧠 Knowledge Base", payload={}),
        cl.Action(name="cmd_machinery_status", label="🔧 Machinery", payload={}),
    ]
    await cl.Message(
        content="## 📋 Quick Commands\n\nChoose an action or type directly in chat:\n",
        actions=all_cmds,
        author="Alex 🤖",
    ).send()


# ── Quick Command Handlers ─────────────────────────────────────────────────
# Each command callback sends a prompt as if the user typed it.

_CMD_MAP = {
    "cmd_renewals":        "Show policies expiring in 30 days and generate renewal emails",
    "cmd_full_report":     "Generate the full report: renewals, cross-sell, compliance, claims",
    "cmd_cross_sell":      "Analyze cross-sell opportunities for all clients",
    "cmd_compliance":      "Check ASF and BaFin compliance",
    "cmd_claims":          "Show open claims and generate follow-up",
    "cmd_send_questionnaire": "Check open claims without questionnaires and send the appropriate damage report form (KFZ-Schaden or Haftpflicht) to each client via email and WhatsApp",
    "cmd_check_forms":     "Show the status of all form submissions: how many sent, in progress, completed, overdue. Include client names and reference numbers",
    "cmd_form_reminders":  "Run form follow-up: check all incomplete form submissions and send reminders to clients who haven't completed their forms after 4, 8, or 12 days",
    "cmd_form_report":     "Generate the daily form completion report: show total forms, completion rate, overdue forms, and recent submissions. Send the report by email to the operator",
    "cmd_check_rca":       "🔍 Type the plate number (e.g. B123ABC) in chat and I will check RCA.",
    "cmd_search_client":   "Search all active clients and show portfolio details",
    "cmd_generate_offers": "Generate offers for clients with expiring policies",
    "cmd_send_emails":     "Generate and queue renewal and cross-sell emails",
    "cmd_daily_brief":     "Generate the daily briefing with all of today's priorities",
    "cmd_oracle_status":   "Connect to Oracle Database and show the full dashboard: connection status, total clients, active policies, monthly premium revenue, open claims with details, and form completion rate. Show all data from Oracle.",
    "cmd_knowledge_status": "Show me what you've learned so far: how many knowledge entries, which hooks are most active, what patterns have you detected, and what are the most used learnings. Also show the Knowledge Base dashboard link: /dashboard/knowledge",
    "cmd_machinery_status": "Show the machinery breakdown claims status: how many equipment entries are registered, how many have active insurance, how many machinery claims are open, investigating, settled, or rejected. List each active claim with its step progress. Also show the Machinery dashboard link: /dashboard/machinery",
}

# Commands that should be sent as user messages (auto-execute)
_CMD_AUTO_EXEC = {"cmd_renewals", "cmd_full_report", "cmd_cross_sell", "cmd_compliance",
                  "cmd_claims", "cmd_search_client", "cmd_generate_offers", "cmd_send_emails",
                  "cmd_daily_brief", "cmd_send_questionnaire", "cmd_check_forms", "cmd_form_reminders",
                  "cmd_oracle_status", "cmd_knowledge_status", "cmd_machinery_status"}

async def _handle_quick_cmd(action: cl.Action):
    """Generic handler for quick commands."""
    cmd_text = _CMD_MAP.get(action.name, "")
    if not cmd_text:
        return

    if action.name in _CMD_AUTO_EXEC:
        # Auto-execute: create fake user message and process through main handler
        await action.remove()
        fake_msg = cl.Message(content=cmd_text, author="user", type="user_message")
        await fake_msg.send()
        # Process through the main handler
        try:
            _on_msg = globals().get("on_message")
            if _on_msg:
                await _on_msg(fake_msg)
            else:
                await cl.Message(content="⚠️ Main handler is not available. Type the command directly in chat.", author="Alex 🤖").send()
        except Exception as _cmd_err:
            import traceback
            await cl.Message(content=f"⚠️ Error executing command: {_cmd_err}", author="Alex 🤖").send()
            traceback.print_exc()
    else:
        # Just display instruction (e.g., check RCA — user needs to type the plate)
        await action.remove()
        await cl.Message(content=cmd_text, author="Alex 🤖").send()

for _cmd_name in _CMD_MAP:
    cl.action_callback(_cmd_name)(_handle_quick_cmd)


@cl.action_callback("agent_local_toggle")
async def _on_agent_toggle(action: cl.Action):
    """Show agent local status and instructions to start/stop."""
    online = action.payload.get("online", False)
    if online:
        # Show status
        from datetime import datetime as _dt3
        info_lines = ["## 🟢 Local Agent — Connected\n"]
        for _ag_id, _ag_info in _cu_agents.items():
            try:
                _last = _dt3.fromisoformat(_ag_info["last_seen"])
                secs = int((_dt3.utcnow() - _last).total_seconds())
                if secs < 120:
                    info_lines.append(f"- **Agent:** `{_ag_info.get('agent_id', _ag_id)}`")
                    info_lines.append(f"- **Platform:** {_ag_info.get('platform', 'N/A')}")
                    conns = _ag_info.get("connectors", [])
                    info_lines.append(f"- **Connectors:** {', '.join(conns) if conns else 'none'}")
                    info_lines.append(f"- **Last signal:** {secs}s ago")
            except Exception:
                pass
        info_lines.append("\n*Local agent is running. You can use: RCA check, internal network access, desktop apps.*")
        await cl.Message(content="\n".join(info_lines), author="Alex 🤖").send()
    else:
        await cl.Message(
            content=(
                "## 🔴 Local Agent — Disconnected\n\n"
                "The local agent is not running. To activate it:\n\n"
                "```bash\n"
                "cd alex-local-agent\n"
                "python main.py start\n"
                "```\n\n"
                "The local agent allows you to:\n"
                "- 🔍 RCA check via CEDAM (ASF portal)\n"
                "- 🖥️ Access desktop apps (Word, Excel, etc.)\n"
                "- 🌐 Access internal network / intranet\n"
                "- 📸 Screenshot and local browser automation"
            ),
            author="Alex 🤖",
        ).send()


@cl.action_callback("open_history")
async def on_open_history(action: cl.Action):
    """Show client-based conversation history."""
    await action.remove()
    meta    = cl.user_session.get("user_meta", {})
    user_id = meta.get("user_id")
    await _show_client_history_picker(user_id)


@cl.action_callback("select_client_history")
async def on_select_client_history(action: cl.Action):
    """User picked a client — show their saved conversations."""
    await action.remove()
    meta        = cl.user_session.get("user_meta", {})
    user_id     = meta.get("user_id")
    client_id   = action.payload["client_id"]
    client_name = action.payload["client_name"]
    await _show_conversation_picker_for_client(user_id, client_id, client_name)


@cl.action_callback("new_conversation_for_client")
async def on_new_conversation_for_client(action: cl.Action):
    """Start a new saved conversation linked to a specific client."""
    await action.remove()
    meta       = cl.user_session.get("user_meta", {})
    user_id    = meta.get("user_id")
    company_id = meta.get("company_id")
    client_id  = action.payload["client_id"]
    client_name = action.payload["client_name"]

    # Create a project for the client if it doesn't exist yet (one project per client)
    project_id = None
    if ADMIN_ENABLED and user_id and company_id:
        # Try to find or create a project named after the client
        projects = list_projects(user_id)
        existing = next((p for p in projects if p["name"] == client_name), None)
        if existing:
            project_id = existing["id"]
        else:
            try:
                proj = create_project(user_id, company_id, client_name,
                                      description=f"Conversations about client {client_name}")
                project_id = proj["id"]
            except Exception:
                project_id = None

        conv = create_conversation(user_id, project_id, title=f"New conversation — {client_name}")
        cl.user_session.set(_SK_CONV_ID,    conv["id"])
        cl.user_session.set(_SK_PROJECT_ID, project_id)
        cl.user_session.set(_SK_CLIENT_ID,  client_id)
        # Link conversation to client immediately
        set_conversation_client(conv["id"], client_id)

    cl.user_session.set(_SK_HISTORY,   [])
    cl.user_session.set(_SK_TITLE_SET, False)
    await cl.Message(
        content=f"💬 New conversation about **{client_name}**. This will be saved automatically.",
        author="Alex 🤖",
    ).send()


@cl.action_callback("new_conversation")
async def on_new_conversation(action: cl.Action):
    await action.remove()
    meta       = cl.user_session.get("user_meta", {})
    user_id    = meta.get("user_id")
    project_id = cl.user_session.get(_SK_PROJECT_ID)

    if ADMIN_ENABLED and user_id and project_id:
        conv = create_conversation(user_id, project_id)
        cl.user_session.set(_SK_CONV_ID, conv["id"])

    cl.user_session.set(_SK_HISTORY,   [])
    cl.user_session.set(_SK_TITLE_SET, False)
    await _show_dashboard_welcome(user_id=user_id, full_name=meta.get("full_name"))


@cl.action_callback("resume_conversation")
async def on_resume_conversation(action: cl.Action):
    await action.remove()
    conv_id = action.payload["conversation_id"]
    title   = action.payload["title"]

    history = load_conversation_history(conv_id) if ADMIN_ENABLED else []

    cl.user_session.set(_SK_CONV_ID,    conv_id)
    cl.user_session.set(_SK_HISTORY,    history)
    cl.user_session.set(_SK_TITLE_SET,  True)
    cl.user_session.set(_SK_SAVE_NUDGE, True)   # already saved — no nudge needed

    if history:
        await _render_history_to_ui(history)

    await cl.Message(
        content=f'📂 Resumed **"{title}"** — {len(history)//2} message(s). Continue below.',
        author="Alex 🤖",
    ).send()


@cl.action_callback("save_conv_pick_client")
async def on_save_conv_pick_client(action: cl.Action):
    """Show client picker so broker can link the current conversation to a client."""
    await action.remove()
    conv_id = cl.user_session.get(_SK_CONV_ID)
    if not conv_id:
        await cl.Message(content="No active conversation to save.", author="Alex 🤖").send()
        return

    clients = get_all_clients_for_picker()
    if not clients:
        await cl.Message(content="No clients in database yet.", author="Alex 🤖").send()
        return

    actions = [
        cl.Action(
            name="save_conv_confirm",
            label=f"👤 {c['name']}",
            payload={"client_id": c["id"], "client_name": c["name"], "conv_id": conv_id},
        )
        for c in clients[:20]
    ]
    await cl.Message(
        content="**Link this conversation to a client:**",
        actions=actions,
        author="Alex 🤖",
    ).send()


@cl.action_callback("save_conv_confirm")
async def on_save_conv_confirm(action: cl.Action):
    """Link the current conversation to the selected client."""
    await action.remove()
    client_id   = action.payload["client_id"]
    client_name = action.payload["client_name"]
    conv_id     = action.payload["conv_id"]

    try:
        set_conversation_client(conv_id, client_id)
        cl.user_session.set(_SK_CLIENT_ID, client_id)

        history = cl.user_session.get("history", [])
        current_title = cl.user_session.get("title_set")
        if not current_title:
            update_conversation_title(conv_id, f"Conversation — {client_name}")

        meta    = cl.user_session.get("user_meta", {})
        user_id = meta.get("user_id")
        elements = []
        if user_id and ADMIN_ENABLED:
            sidebar_content = await _build_sidebar_content(user_id)
            elements.append(cl.Text(
                name="📁 Saved Conversations",
                content=sidebar_content,
                display="side",
            ))
        await cl.Message(
            content=f"✅ Conversation saved and linked to **{client_name}**.\n"
                    f"You can find it in the side panel →",
            elements=elements,
            author="Alex 🤖",
        ).send()
    except Exception as e:
        await cl.Message(content=f"Could not link conversation: {e}", author="Alex 🤖").send()


# ── Search & Delete Conversations ─────────────────────────────────────────────

@cl.action_callback("search_conversations")
async def on_search_conversations(action: cl.Action):
    """Prompt user for search query, then show matching conversations."""
    await action.remove()
    meta    = cl.user_session.get("user_meta", {})
    user_id = meta.get("user_id")
    if not user_id or not ADMIN_ENABLED:
        await cl.Message(content="Feature unavailable.", author="Alex 🤖").send()
        return

    # Ask user what to search for
    res = await cl.AskUserMessage(
        content="🔍 **Search conversations** — type your keywords:",
        author="Alex 🤖",
    ).send()
    if not res:
        return

    query = res["output"] if isinstance(res, dict) else res.content
    results = search_conversations(user_id, query.strip())

    if not results:
        await cl.Message(
            content=f"No conversations found containing **\"{query}\"**.",
            actions=[cl.Action(name="search_conversations", label="🔍 Search again", payload={})],
            author="Alex 🤖",
        ).send()
        return

    actions = []
    lines = [f"## 🔍 Results: \"{query}\" ({len(results)})\n"]
    for conv in results[:15]:
        title = conv.get("title", "Conversation")
        msgs  = conv.get("message_count", 0)
        client = conv.get("client_name") or "No client linked"
        upd   = conv.get("updated_at", "")[:10]
        lines.append(f"- **{title}** — {client} ({msgs} msg, {upd})")
        actions.append(cl.Action(
            name="resume_conversation",
            label=f"📂 {title[:30]}",
            payload={"conversation_id": conv["id"], "title": title},
        ))
        actions.append(cl.Action(
            name="delete_conversation_confirm",
            label=f"🗑️ Delete: {title[:25]}",
            payload={"conversation_id": conv["id"], "title": title},
        ))

    actions.append(cl.Action(name="search_conversations", label="🔍 Search again", payload={}))

    await cl.Message(
        content="\n".join(lines),
        actions=actions,
        author="Alex 🤖",
    ).send()


@cl.action_callback("delete_conversation_confirm")
async def on_delete_conversation_confirm(action: cl.Action):
    """Delete a saved conversation after confirmation."""
    await action.remove()
    meta    = cl.user_session.get("user_meta", {})
    user_id = meta.get("user_id")
    conv_id = action.payload["conversation_id"]
    title   = action.payload["title"]

    if not user_id or not ADMIN_ENABLED:
        await cl.Message(content="Feature unavailable.", author="Alex 🤖").send()
        return

    # Confirm deletion
    actions = [
        cl.Action(name="delete_conversation_do", label="✅ Yes, delete",
                  payload={"conversation_id": conv_id, "title": title}),
        cl.Action(name="delete_conversation_cancel", label="❌ Cancel", payload={}),
    ]
    await cl.Message(
        content=f"⚠️ Are you sure you want to delete the conversation **\"{title}\"**?\n\nThis action is irreversible.",
        actions=actions,
        author="Alex 🤖",
    ).send()


@cl.action_callback("delete_conversation_do")
async def on_delete_conversation_do(action: cl.Action):
    """Actually delete the conversation."""
    await action.remove()
    meta    = cl.user_session.get("user_meta", {})
    user_id = meta.get("user_id")
    conv_id = action.payload["conversation_id"]
    title   = action.payload["title"]

    success = delete_conversation(conv_id, user_id) if ADMIN_ENABLED else False
    if success:
        await cl.Message(
            content=f"🗑️ Conversation **\"{title}\"** has been deleted.",
            author="Alex 🤖",
        ).send()
    else:
        await cl.Message(
            content=f"❌ Could not delete the conversation. Check if it exists.",
            author="Alex 🤖",
        ).send()


@cl.action_callback("delete_conversation_cancel")
async def on_delete_conversation_cancel(action: cl.Action):
    """Cancel deletion."""
    await action.remove()
    await cl.Message(content="Cancelled — the conversation was not deleted.", author="Alex 🤖").send()


# ── Output file cleanup — interactive flow ────────────────────────────────────

@cl.action_callback("output_cleanup_start")
async def on_output_cleanup_start(action: cl.Action):
    """Step 1: Ask what to keep — show age/type filters."""
    await action.remove()
    from datetime import datetime as _dt
    _files = sorted(
        [f for f in OUTPUT_DIR.iterdir() if f.is_file() and f.suffix.lower() in {".pdf", ".txt", ".xlsx", ".docx"}],
        key=lambda f: f.stat().st_mtime, reverse=True
    )
    if not _files:
        await cl.Message(content="📂 No files to delete.", author="Alex 🤖").send()
        return

    # Categorize by age
    _now = _dt.now().timestamp()
    _old_30  = [f for f in _files if (_now - f.stat().st_mtime) > 30 * 86400]
    _old_7   = [f for f in _files if 7 * 86400 < (_now - f.stat().st_mtime) <= 30 * 86400]
    _recent  = [f for f in _files if (_now - f.stat().st_mtime) <= 7 * 86400]
    _total_old_mb = sum(f.stat().st_size for f in _old_30) / (1024 * 1024)

    summary = (
        f"**📂 {len(_files)} files** in total:\n"
        f"- 🟢 Recent (<7 days): **{len(_recent)}** files\n"
        f"- 🟡 7-30 days: **{len(_old_7)}** files\n"
        f"- 🔴 Old (>30 days): **{len(_old_30)}** files ({_total_old_mb:.1f} MB)\n\n"
        f"What do you want to delete?"
    )

    await cl.Message(
        content=summary,
        actions=[
            cl.Action(name="output_cleanup_confirm", label=f"🔴 Delete >30 days ({len(_old_30)} files, {_total_old_mb:.1f} MB)",
                      payload={"mode": "old30"}),
            cl.Action(name="output_cleanup_confirm", label=f"🟡 Delete >7 days ({len(_old_7) + len(_old_30)} files)",
                      payload={"mode": "old7"}),
            cl.Action(name="output_cleanup_confirm", label=f"🗑️ Delete ALL ({len(_files)} files)",
                      payload={"mode": "all"}),
            cl.Action(name="output_cleanup_cancel", label="❌ Cancel",
                      payload={"mode": "cancel"}),
        ],
        author="Alex 🤖",
    ).send()


@cl.action_callback("output_cleanup_confirm")
async def on_output_cleanup_confirm(action: cl.Action):
    """Step 2: Show exact list of files to be deleted — ask final confirmation."""
    await action.remove()
    mode = action.payload.get("mode", "cancel")

    from datetime import datetime as _dt
    _files = sorted(
        [f for f in OUTPUT_DIR.iterdir() if f.is_file() and f.suffix.lower() in {".pdf", ".txt", ".xlsx", ".docx"}],
        key=lambda f: f.stat().st_mtime, reverse=True
    )
    _now = _dt.now().timestamp()

    if mode == "old30":
        _to_delete = [f for f in _files if (_now - f.stat().st_mtime) > 30 * 86400]
        _keep = [f for f in _files if (_now - f.stat().st_mtime) <= 30 * 86400]
    elif mode == "old7":
        _to_delete = [f for f in _files if (_now - f.stat().st_mtime) > 7 * 86400]
        _keep = [f for f in _files if (_now - f.stat().st_mtime) <= 7 * 86400]
    elif mode == "all":
        _to_delete = list(_files)
        _keep = []
    else:
        await cl.Message(content="✅ Cancelled — no files deleted.", author="Alex 🤖").send()
        return

    if not _to_delete:
        await cl.Message(content="✅ No files match the criteria.", author="Alex 🤖").send()
        return

    # Store list in session for final delete step
    cl.user_session.set("_cleanup_pending", [str(f) for f in _to_delete])

    _del_mb = sum(f.stat().st_size for f in _to_delete) / (1024 * 1024)
    _lines = [f"### ⚠️ {len(_to_delete)} files will be deleted ({_del_mb:.1f} MB):\n"]
    for f in _to_delete[:20]:
        _mtime = _dt.fromtimestamp(f.stat().st_mtime).strftime("%d %b %Y")
        _sz = f"{f.stat().st_size/1024:.0f} KB"
        _lines.append(f"- `{f.name}` — {_sz} — {_mtime}")
    if len(_to_delete) > 20:
        _lines.append(f"- _... and {len(_to_delete) - 20} more files_")

    if _keep:
        _lines.append(f"\n### ✅ {len(_keep)} newer files will be kept:")
        for f in _keep[:5]:
            _lines.append(f"- `{f.name}`")
        if len(_keep) > 5:
            _lines.append(f"  _... and {len(_keep) - 5} others_")

    _lines.append("\n**Confirm deletion?** This action cannot be undone.")

    await cl.Message(
        content="\n".join(_lines),
        actions=[
            cl.Action(name="output_cleanup_execute", label=f"✅ Yes, delete {len(_to_delete)} files",
                      payload={"confirmed": True}),
            cl.Action(name="output_cleanup_cancel", label="❌ No, cancel",
                      payload={}),
        ],
        author="Alex 🤖",
    ).send()


@cl.action_callback("output_cleanup_execute")
async def on_output_cleanup_execute(action: cl.Action):
    """Step 3: Execute the confirmed deletion."""
    await action.remove()
    _pending = cl.user_session.get("_cleanup_pending", [])
    if not _pending:
        await cl.Message(content="⚠️ Nothing to delete.", author="Alex 🤖").send()
        return

    _deleted = []
    _errors = []
    for _path_str in _pending:
        try:
            _p = Path(_path_str)
            if _p.exists():
                _p.unlink()
                _deleted.append(_p.name)
        except Exception as _e:
            _errors.append(f"{Path(_path_str).name}: {_e}")

    cl.user_session.set("_cleanup_pending", [])

    _remaining = [f for f in OUTPUT_DIR.iterdir() if f.is_file() and f.suffix.lower() in {".pdf", ".txt", ".xlsx", ".docx"}]
    _lines = [f"### ✅ Cleanup completed\n"]
    _lines.append(f"- **Deleted:** {len(_deleted)} files")
    if _errors:
        _lines.append(f"- **Errors:** {len(_errors)}")
        for _e in _errors[:3]:
            _lines.append(f"  - {_e}")
    _lines.append(f"- **Remaining:** {len(_remaining)} files in output/")
    if _deleted[:5]:
        _lines.append(f"\n_Deleted examples: {', '.join(_deleted[:5])}_")

    await cl.Message(content="\n".join(_lines), author="Alex 🤖").send()


@cl.action_callback("output_cleanup_cancel")
async def on_output_cleanup_cancel(action: cl.Action):
    """Cancel cleanup."""
    await action.remove()
    cl.user_session.set("_cleanup_pending", [])
    await cl.Message(content="✅ Cancelled — no files deleted.", author="Alex 🤖").send()


# ── Offer action callbacks (persistent buttons — no timeout) ─────────────────

@cl.action_callback("offer_approve")
async def on_offer_approve(action: cl.Action):
    """Broker approved offer → inject approval into history and trigger agentic loop."""
    await action.remove()
    base_title = cl.user_session.get("last_offer_title", "")

    import re as _re
    _m = _re.search(r"(OFF-[A-F0-9]{8})", base_title)
    _offer_id_hint = _m.group(1) if _m else ""

    cl.user_session.set("offer_approved", True)

    approval_text = (
        f"Send offer {_offer_id_hint} via email to the client."
        if _offer_id_hint else
        "Send the offer via email to the client."
    )

    await cl.Message(
        content="\u2705 **Offer approved!** Preparing to send via email...",
        author="Alex \U0001f916"
    ).send()

    # Feed this through the on_message handler directly
    synthetic = cl.Message(content=approval_text, author="user")
    await on_message(synthetic)


@cl.action_callback("offer_save_dashboard")
async def on_offer_save_dashboard(action: cl.Action):
    """Save offer to approval_queue dashboard."""
    await action.remove()
    offer_content = cl.user_session.get("last_offer_content", "")
    base_title = cl.user_session.get("last_offer_title", "")
    history = cl.user_session.get("history", [])
    _save_msg = ""
    try:
        import uuid as _uuid_dash, sqlite3 as _sql_dash
        from pathlib import Path as _P_dash
        import re as _re_dash
        _db_path = str(_P_dash(__file__).parent / "mcp-server" / "insurance_broker.db")
        _conn_dash = _sql_dash.connect(_db_path)
        _m_offer = _re_dash.search(r"(OFF-[A-F0-9]{8})", base_title)
        _offer_id_d = _m_offer.group(1) if _m_offer else ""
        _client_name_d = base_title.split("_")[0] if "_" in base_title else "Client"
        _client_id_d = ""
        _client_email_d = ""
        try:
            if _offer_id_d:
                _offer_row = _conn_dash.execute(
                    "SELECT o.client_id, c.name, c.email FROM offers o LEFT JOIN clients c ON c.id = o.client_id WHERE o.id = ?",
                    (_offer_id_d,)
                ).fetchone()
            else:
                _offer_row = None
            if _offer_row:
                _client_id_d = _offer_row[0] or ""
                _client_name_d = _offer_row[1] or _client_name_d
                _client_email_d = _offer_row[2] or ""
        except Exception:
            pass
        if not _client_id_d:
            _client_id_d = "unknown"
        _approval_id = _uuid_dash.uuid4().hex[:16]
        _conn_dash.execute(
            "INSERT OR IGNORE INTO approval_queue "
            "(id, type, client_id, client_name, client_email, subject, email_body_html, offer_id, status, created_at) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'pending', datetime('now'))",
            (_approval_id, "offer", _client_id_d, _client_name_d, _client_email_d,
             f"Offer {_offer_id_d or 'new'} - {_client_name_d}",
             offer_content or "Offer generated by Alex",
             _offer_id_d or "")
        )
        _conn_dash.commit()
        _conn_dash.close()
        _save_msg = (
            f"\u2705 **Offer saved to Dashboard!**\n\n"
            f"- **Client:** {_client_name_d}\n"
            f"- **Offer:** {_offer_id_d or 'N/A'}\n"
            f"- **Status:** Pending approval\n\n"
            f"\U0001f449 [Open Dashboard \u2192](/dashboard/approvals)"
        )
    except Exception as _e_dash:
        import traceback as _tb_dash
        _tb_dash.print_exc()
        _save_msg = f"\u274c Error saving to dashboard: {_e_dash}"

    history.append({"role": "assistant", "content": _save_msg or "Offer saved."})
    cl.user_session.set("history", history)
    await cl.Message(content=_save_msg, author="Alex \U0001f916").send()


@cl.action_callback("offer_download")
async def on_offer_download(action: cl.Action):
    """Export offer as PDF/XLSX/DOCX."""
    await action.remove()
    offer_content = cl.user_session.get("last_offer_content", "")
    base_title = cl.user_session.get("last_offer_title", "")
    history = cl.user_session.get("history", [])

    history.append({"role": "assistant", "content": "Offer downloaded."})
    cl.user_session.set("history", history)
    await send_export_files(offer_content, base_title)


@cl.action_callback("offer_edit")
async def on_offer_edit(action: cl.Action):
    """Broker wants to edit the offer."""
    await action.remove()
    history = cl.user_session.get("history", [])
    history.append({"role": "assistant", "content": "Broker wants to edit the offer."})
    cl.user_session.set("history", history)
    await cl.Message(
        content=(
            "\u270f\ufe0f **What would you like to change?** Write naturally, for example:\n"
            "- *'Change validity to 14 days'*\n"
            "- *'Add a note about 10% discount'*\n"
            "- *'Generate in Romanian'*\n"
            "- *'Remove the Generali product'*"
        ),
        author="Alex \U0001f916"
    ).send()


# ─────────────────────────────────────────────────────────────────────────────

@cl.on_message
async def on_message(message: cl.Message):
    """Handle incoming broker message — full agentic loop with Claude."""
    history = cl.user_session.get("history", [])

    # ── History sanitizer: fix orphan tool_use without tool_result ─────────
    # This prevents 400 errors when a previous flow returned early
    if history:
        _sanitized = []
        _pending_tool_ids = []
        for _msg in history:
            if _msg.get("role") == "assistant":
                content = _msg.get("content", [])
                if isinstance(content, list):
                    for _block in content:
                        if hasattr(_block, "type") and _block.type == "tool_use":
                            _pending_tool_ids.append(_block.id)
                        elif isinstance(_block, dict) and _block.get("type") == "tool_use":
                            _pending_tool_ids.append(_block.get("id"))
                _sanitized.append(_msg)
            elif _msg.get("role") == "user":
                content = _msg.get("content", [])
                if isinstance(content, list):
                    for _block in content:
                        if isinstance(_block, dict) and _block.get("type") == "tool_result":
                            _tid = _block.get("tool_use_id")
                            if _tid in _pending_tool_ids:
                                _pending_tool_ids.remove(_tid)
                _sanitized.append(_msg)
            else:
                _sanitized.append(_msg)
        # If there are orphan tool_use IDs, add fake tool_results
        if _pending_tool_ids:
            _fake_results = [{
                "type": "tool_result",
                "tool_use_id": _tid,
                "content": "[Session interrupted — result unavailable]",
            } for _tid in _pending_tool_ids]
            _sanitized.append({"role": "user", "content": _fake_results})
            history = _sanitized

    # ── Project / conversation persistence ────────────────────────────────
    conv_id    = cl.user_session.get(_SK_CONV_ID)
    project_id = cl.user_session.get(_SK_PROJECT_ID)
    client_id  = cl.user_session.get(_SK_CLIENT_ID)
    title_set  = cl.user_session.get(_SK_TITLE_SET, False)

    # Auto-create a conversation on the first message (so nothing is lost)
    if ADMIN_ENABLED and not conv_id and message.content:
        meta       = cl.user_session.get("user_meta", {})
        _user_id   = meta.get("user_id")
        _company_id = meta.get("company_id")
        if _user_id and _company_id:
            try:
                conv = create_conversation(_user_id, project_id)
                conv_id = conv["id"]
                cl.user_session.set(_SK_CONV_ID, conv_id)
            except Exception:
                pass  # non-fatal — conversation won't be saved but chat works

    # Auto-title from the first user message — clean, max 60 chars, no markdown
    if conv_id and not title_set and message.content:
        _raw = message.content.strip().replace("\n", " ")
        # Strip common prefixes brokers type
        import re as _re2
        _raw = _re2.sub(r'^(vreau|vrea|as vrea|as dori|cauta|caută|arata|arată|fa|fă|verifica|verifică|spune|hai|hei|alex[,:]?\s*)', '', _raw, flags=_re2.IGNORECASE).strip()
        _title = (_raw[:57] + "…") if len(_raw) > 60 else _raw
        _title = _title or "Conversation"
        update_conversation_title(conv_id, _title)
        cl.user_session.set(_SK_TITLE_SET, True)
    # ─────────────────────────────────────────────────────────────────────

    # ── Sidebar shortcut — intercept starter message before Claude sees it ──
    _msg_lower = (message.content or "").strip().lower()
    if _msg_lower in ("arată-mi conversațiile salvate", "show saved conversations", "saved conversations") or "conversații salvate" in _msg_lower or "saved conversations" in _msg_lower:
        _meta = cl.user_session.get("user_meta", {})
        _uid  = _meta.get("user_id")
        if _uid and ADMIN_ENABLED:
            _sidebar_content = await _build_sidebar_content(_uid)
            _elements = [cl.Text(
                name="📁 Saved Conversations",
                content=_sidebar_content,
                display="side",
            )]
            await cl.Message(
                content="📁 **Saved Conversations** — the panel has opened on the right.",
                elements=_elements,
                author="Alex 🤖",
            ).send()
            return
    # ────────────────────────────────────────────────────────────────────────

    # Build user message content (text + any uploaded files)
    user_content = []
    if message.content and message.content.strip():
        user_content.append({"type": "text", "text": message.content})

    # Process any uploaded files inline
    if message.elements:
        for element in message.elements:
            has_data = (getattr(element, "content", None) or getattr(element, "path", None))
            if has_data and hasattr(element, "name"):
                try:
                    async with cl.Step(name=f"📄 Analyzing {element.name}...", type="tool", show_input=False) as step:
                        step.output = "Processing with Claude Vision..."
                        analysis = await process_uploaded_file(element)

                    # Show extracted data with confirmation buttons
                    await cl.Message(
                        content=(
                            f"✅ **Document read: {element.name}**\n\n"
                            f"{analysis}\n\n"
                            f"---\n"
                            f"⚠️ **Check the data above before continuing.** "
                            f"If something is wrong (e.g. clinic phone confused with client phone), "
                            f"type your correction in the next message."
                        ),
                        author="Alex 🤖"
                    ).send()

                    # Ask broker to confirm before proceeding
                    res = await cl.AskActionMessage(
                        content="Is the data correct?",
                        actions=[
                            cl.Action(name="confirm", label="✅ Correct — continue", payload={"value": "confirm"}),
                            cl.Action(name="edit", label="✏️ I want to correct something", payload={"value": "edit"}),
                        ],
                        author="Alex 🤖",
                        timeout=60,
                    ).send()

                    if res and res.get("payload", {}).get("value") == "edit":
                        await cl.Message(
                            content="✏️ Type your correction (e.g. 'The client name is Ion Popescu, the phone is 0722111222, not the clinic phone').",
                            author="Alex 🤖"
                        ).send()
                        # Add analysis to context but flag that broker will correct it
                        user_content.append({"type": "text", "text": (
                            f"[Document analizat: {element.name}]\n{analysis}\n\n"
                            f"[IMPORTANT: Brokerul va corecta datele în mesajul următor. Asteaptă corecția înainte să cauți sau să creezi clientul.]"
                        )})
                    else:
                        # Confirmed or timeout — proceed with analysis
                        user_content.append({"type": "text", "text": analysis})

                except Exception as e:
                    await cl.Message(
                        content=f"⚠️ Could not process **{element.name}**: {str(e)[:200]}",
                        author="Alex 🤖"
                    ).send()

    if not user_content:
        return

    # ── Shortcut: "show offer" without calling Claude ─────────────────────────
    msg_lower = (message.content or "").lower().strip()
    SHOW_OFFER_TRIGGERS = [
        "sa vad oferta", "să văd oferta", "show offer", "arata oferta", "arată oferta",
        "show me the offer", "view offer", "see offer", "oferta", "show the offer",
        "vreau sa vad", "vreau să văd",
    ]
    if any(t in msg_lower for t in SHOW_OFFER_TRIGGERS):
        last_file = cl.user_session.get("last_offer_file")
        last_content = cl.user_session.get("last_offer_content")
        if last_file and Path(last_file).exists():
            fname = Path(last_file).name
            await cl.Message(
                content=f"📄 **Here is the last generated offer:**",
                elements=[cl.File(name=fname, path=last_file, display="inline")],
                author="Alex 🤖"
            ).send()
            # Also re-send exports
            last_title = cl.user_session.get("last_offer_title", "Offer")
            await send_export_files(last_content, last_title)
            return
        # No cached offer — let Claude handle it normally
    # ─────────────────────────────────────────────────────────────────────────

    history.append({"role": "user", "content": user_content})

    # ── ORCHESTRATOR: classify intent and route to specialist module ──────
    _orch_msg_text = message.content or ""
    _orch_lower = _orch_msg_text.lower()

    # Intent classification rules (keyword-based, fast, no API call)
    _ORCHESTRATOR_MODULES = [
        {
            "id": "claims",
            "name": "🔴 Claims Module",
            "icon": "🔴",
            "keywords": ["daună", "daune", "claim", "claims", "schadenm", "schaden", "accident",
                         "pagubă", "despăgubire", "constatare", "reparație", "damage", "incident"],
            "description": "Damage registration, claim tracking, status updates",
        },
        {
            "id": "forms",
            "name": "📋 Forms Module",
            "icon": "📋",
            "keywords": ["formular", "chestionar", "questionnaire", "form", "fragebogen",
                         "completare", "link", "trimite link", "send link", "reminder",
                         "incomplete", "follow-up", "submission"],
            "description": "Questionnaire management, send links, track completions",
        },
        {
            "id": "clients",
            "name": "👤 Client Module",
            "icon": "👤",
            "keywords": ["client", "clienți", "kunde", "kunden", "caută client", "search client",
                         "adaugă client", "add client", "profil", "profile", "contact",
                         "telefon", "email client", "portofoliu"],
            "description": "Client search, profiles, portfolio management",
        },
        {
            "id": "policies",
            "name": "📄 Policy Module",
            "icon": "📄",
            "keywords": ["poliță", "polița", "polite", "policy", "policies", "rca", "casco",
                         "reînnoire", "renewal", "expiră", "expired", "asigurare", "insurance",
                         "kfz", "haftpflicht", "versicherung", "pad", "home"],
            "description": "Policy lookup, renewals, coverage verification",
        },
        {
            "id": "offers",
            "name": "💰 Offers Module",
            "icon": "💰",
            "keywords": ["ofertă", "oferta", "offer", "compare", "compară", "preț", "price",
                         "premium", "angebot", "cotație", "quote", "cheapest", "ieftin"],
            "description": "Price comparison, offer generation, quotes",
        },
        {
            "id": "compliance",
            "name": "⚖️ Compliance Module",
            "icon": "⚖️",
            "keywords": ["asf", "bafin", "compliance", "raport lunar", "monthly report",
                         "regulatory", "conformitate", "audit"],
            "description": "ASF/BaFin reports, regulatory compliance checks",
        },
        {
            "id": "vehicles",
            "name": "🚗 Vehicle Module",
            "icon": "🚗",
            "keywords": ["vehicul", "mașină", "auto", "vehicle", "car", "fahrzeug",
                         "înmatriculare", "plate", "vin", "serie sasiu", "motor"],
            "description": "Vehicle management, registration, fleet tracking",
        },
        {
            "id": "knowledge",
            "name": "🧠 Knowledge Base (RAG)",
            "icon": "🧠",
            "keywords": ["ce acoperă", "what covers", "exclusion", "excludere",
                         "documente necesare", "documents needed", "cum funcționează",
                         "how does", "knowledge", "explică", "explain"],
            "description": "Semantic search in product docs, guides, FAQs",
        },
        {
            "id": "reports",
            "name": "📊 Reports Module",
            "icon": "📊",
            "keywords": ["raport", "report", "statistică", "statistics", "dashboard",
                         "kpi", "briefing", "daily", "zilnic", "rezumat", "summary"],
            "description": "Daily briefings, KPIs, analytics, reports",
        },
        {
            "id": "automation",
            "name": "🤖 Automation Module",
            "icon": "🤖",
            "keywords": ["cron", "automat", "automatic", "schedule", "programează",
                         "reminder", "task", "job", "workflow"],
            "description": "Cron jobs, scheduled tasks, automated workflows",
        },
        {
            "id": "documents",
            "name": "📎 Document Module",
            "icon": "📎",
            "keywords": ["upload", "document", "pdf", "scan", "poză", "photo", "image",
                         "fișier", "file", "drive", "sharepoint"],
            "description": "Document upload, AI analysis, cloud storage sync",
        },
        {
            "id": "oracle",
            "name": "🗄️ Oracle DB Module",
            "icon": "🗄️",
            "keywords": ["oracle", "database", "baza de date", "sync", "sincronizare",
                         "sql", "tabel", "table", "query"],
            "description": "Oracle database sync, queries, data management",
        },
    ]

    # Score each module
    _orch_scores = []
    for _mod in _ORCHESTRATOR_MODULES:
        _score = sum(1 for kw in _mod["keywords"] if kw in _orch_lower)
        if _score > 0:
            _orch_scores.append((_mod, _score))

    _orch_scores.sort(key=lambda x: x[1], reverse=True)

    # Determine routing
    if _orch_scores:
        _primary = _orch_scores[0][0]
        _secondary = [s[0] for s in _orch_scores[1:3]]  # up to 2 secondary modules
    else:
        _primary = {"id": "general", "name": "💬 General Assistant", "icon": "💬",
                     "description": "General broker assistance"}
        _secondary = []

    # Show orchestrator routing as a visible step
    _routing_detail = f"**{_primary['name']}** — {_primary['description']}"
    if _secondary:
        _routing_detail += "\n" + "\n".join(
            f"  ↳ {m['name']}" for m in _secondary
        )

    async with cl.Step(name=f"🧠 Orchestrator → {_primary['name']}", type="run", show_input=False) as _orch_step:
        _orch_step.output = _routing_detail
    # ── END ORCHESTRATOR ─────────────────────────────────────────────────────

    final_text = ""
    iterations = 0
    import asyncio as _asyncio

    while iterations < 10:
        iterations += 1

        # Call Claude API (run in thread pool to avoid blocking the event loop)
        response = None
        last_error = None
        for _attempt in range(3):
            try:
                response = await _asyncio.get_event_loop().run_in_executor(
                    None,
                    lambda: client.messages.create(
                        model=MODEL,
                        max_tokens=4096,
                        system=[{
                            "type": "text",
                            "text": SYSTEM_PROMPT + _get_rag_context(message.content if hasattr(message, 'content') else str(message)),
                            "cache_control": {"type": "ephemeral"},
                        }],
                        tools=TOOLS,
                        temperature=0.1,
                        messages=history,
                        extra_headers={"anthropic-beta": "prompt-caching-2024-07-31"},
                    )
                )
                # Track token usage (includes cache_creation_input_tokens + cache_read_input_tokens)
                if ADMIN_ENABLED and response.usage:
                    _meta = cl.user_session.get("user_meta", {})
                    if _meta.get("user_id"):
                        total_tokens = (response.usage.input_tokens or 0) + (response.usage.output_tokens or 0)
                        record_token_usage(
                            company_id=_meta.get("company_id"),
                            user_id=_meta["user_id"],
                            tokens=total_tokens,
                        )

                break  # success
            except Exception as e:
                last_error = str(e)
                err_lower = last_error.lower()
                if "overloaded" in err_lower or "529" in last_error or "503" in last_error or "429" in last_error:
                    wait_sec = [5, 15, 30][min(_attempt, 2)]
                    await _asyncio.sleep(wait_sec)
                    continue
                # Context too long — trim history and retry
                if "too long" in err_lower or "too many" in err_lower or "400" in last_error:
                    if len(history) > 6:
                        history = history[:2] + history[-4:]
                        cl.user_session.set("history", history)
                        continue
                # Hard error — log it so we can debug
                import logging as _logging
                _logging.getLogger("alex").error(f"[on_message] hard error: {last_error}")
                print(f"[alex] HARD ERROR in agentic loop: {last_error}", flush=True)
                await cl.Message(
                    content=f"Something went wrong. Can you rephrase the request or try again?\n\n*(Internal error: {last_error[:200]})*",
                    author="Alex 🤖"
                ).send()
                return

        if response is None:
            await cl.Message(
                content="⚠️ AI service is overloaded at the moment. Please resend your message.",
                author="Alex 🤖"
            ).send()
            return

        # Parse response content blocks
        text_parts = []
        tool_use_blocks = []

        for block in response.content:
            if block.type == "text" and block.text and block.text.strip():
                text_parts.append(block.text)
            elif block.type == "tool_use":
                tool_use_blocks.append(block)

        # Append assistant turn to history
        history.append({"role": "assistant", "content": response.content})

        if not tool_use_blocks:
            final_text = "\n".join(text_parts) if text_parts else ""

            # ── Hallucination guard ────────────────────────────────────────
            # Check if recent history has any tool results (means we already called tools)
            recent_tool_results = [
                msg for msg in history[-4:]
                if msg.get("role") == "user" and
                isinstance(msg.get("content"), list) and
                any(c.get("type") == "tool_result" for c in msg["content"])
            ]
            answer_lower = final_text.lower()
            negative_phrases = [
                "nu am gasit", "nu am găsit", "nu exista", "nu există",
                "not found", "no products", "no results", "cannot find",
                "nu sunt disponibile", "nu s-au gasit", "nu s-au găsit",
                "lipsesc", "indisponibil",
            ]
            has_negative = any(neg in answer_lower for neg in negative_phrases)

            looks_like_hallucination = (
                not recent_tool_results and
                not has_negative and
                any(kw in answer_lower for kw in [
                    "allianz", "generali", "omniasig", "nn asigurari", "groupama",
                    "brd asigurari", "asirom", "uniqa", "signal iduna",
                    "premium anual", "prima anuala", "annual premium",
                    "here are the", "iată produsele", "top health", "top rca",
                    "recommend the", "recomand", "produsele disponibile",
                ]) and
                any(kw in answer_lower for kw in [
                    "health", "rca", "casco", "life", "pad", "cmr", "kfz",
                    "asigurare", "insurance", "polita", "poliță",
                ])
            )
            if looks_like_hallucination:
                # Discard the hallucinated answer and force tool call
                history.pop()  # remove assistant turn with hallucinated text
                history.append({"role": "user", "content": [{
                    "type": "text",
                    "text": (
                        "Nu răspunde din memorie. Apelează imediat broker_search_products "
                        "cu tipul de produs potrivit și arată rezultatele reale din baza de date."
                    )
                }]})
                final_text = ""
                continue  # retry — Claude will now call the tool
            # ────────────────────────────────────────────────────────────────

            # stop_reason == "end_turn" with no tool calls → done
            break

        # ── Execute tool calls ─────────────────────────────────────────────
        tool_results = []

        allowed_tools = cl.user_session.get("allowed_tools")  # None = all allowed
        user_meta = cl.user_session.get("user_meta", {})

        for tb in tool_use_blocks:
            tool_name = tb.name
            tool_input = tb.input if isinstance(tb.input, dict) else {}
            tool_use_id = tb.id

            # ── Permission gate ────────────────────────────────────────────
            if allowed_tools is not None and tool_name not in allowed_tools:
                denied_result = f"⛔ Access denied: you don't have permission to use '{tool_name}'. Contact your admin."
                if ADMIN_ENABLED and user_meta.get("user_id"):
                    log_audit(
                        user_id=user_meta["user_id"],
                        company_id=user_meta.get("company_id"),
                        tool_name=tool_name,
                        input_summary=str(tool_input)[:200],
                        success=False,
                        tokens=0,
                    )
                tool_results.append({
                    "type": "tool_result",
                    "tool_use_id": tool_use_id,
                    "content": denied_result,
                })
                continue

            async with cl.Step(name=f"🔧 {tool_name}", type="tool", show_input=True) as step:
                step.input = json.dumps(tool_input, indent=2, ensure_ascii=False)

                # ── Save conversation → link to client ────────────────────
                if tool_name == "broker_save_conversation":
                    _client_id = tool_input.get("client_id", "").strip()
                    _new_title  = tool_input.get("title", "").strip()
                    _conv_id   = cl.user_session.get(_SK_CONV_ID)

                    if not ADMIN_ENABLED or not _client_id:
                        result = "⚠️ Cannot save: admin features disabled or no client_id provided."
                    elif not _conv_id:
                        # No conversation yet — create one now
                        _meta       = cl.user_session.get("user_meta", {})
                        _uid        = _meta.get("user_id")
                        _cid_comp   = _meta.get("company_id")
                        if _uid and _cid_comp:
                            # Find or create a project for this client
                            _projects  = list_projects(_uid)
                            # Try to get client name for project name
                            import sqlite3 as _sq
                            _db_conn = None
                            _client_name = _client_id
                            try:
                                from shared.db import get_conn as _gcn
                                _db_conn = _gcn()
                                _cr = _db_conn.execute("SELECT name FROM clients WHERE id=?", (_client_id,)).fetchone()
                                if _cr:
                                    _client_name = _cr["name"]
                            except Exception:
                                pass
                            finally:
                                if _db_conn:
                                    _db_conn.close()
                            _existing_proj = next((p for p in _projects if p["name"] == _client_name), None)
                            if _existing_proj:
                                _proj_id = _existing_proj["id"]
                            else:
                                try:
                                    _proj = create_project(_uid, _cid_comp, _client_name)
                                    _proj_id = _proj["id"]
                                except Exception:
                                    _proj_id = None
                            _conv = create_conversation(_uid, _proj_id,
                                                        title=_new_title or f"Conversation about {_client_name}")
                            _conv_id = _conv["id"]
                            cl.user_session.set(_SK_CONV_ID,    _conv_id)
                            cl.user_session.set(_SK_PROJECT_ID, _proj_id)
                            cl.user_session.set(_SK_CLIENT_ID,  _client_id)
                            # Save history so far
                            _hist = cl.user_session.get(_SK_HISTORY, [])
                            if _hist:
                                save_conversation_history(_conv_id, _hist)
                            result = f"✅ Conversation created and linked to **{_client_name}** ({_client_id}). It will appear in '📁 Conversation history by client'."
                        else:
                            result = "⚠️ Cannot save: user session not initialized."
                    else:
                        # Conversation exists — just link it
                        try:
                            set_conversation_client(_conv_id, _client_id)
                            cl.user_session.set(_SK_CLIENT_ID, _client_id)
                            if _new_title:
                                update_conversation_title(_conv_id, _new_title)
                                cl.user_session.set(_SK_TITLE_SET, True)
                            # Get client name for friendly message
                            _client_display = _client_id
                            try:
                                from shared.db import get_conn as _gcn2
                                _dc = _gcn2()
                                _cr2 = _dc.execute("SELECT name FROM clients WHERE id=?", (_client_id,)).fetchone()
                                if _cr2:
                                    _client_display = _cr2["name"]
                                _dc.close()
                            except Exception:
                                pass
                            result = f"✅ Conversation linked to **{_client_display}** ({_client_id}). It will appear in '📁 Conversation history by client'."
                        except Exception as _e:
                            result = f"❌ Could not link conversation: {_e}"

                # ── Async computer-use tools ──────────────────────────────
                elif tool_name == "broker_computer_use_status":
                    result = _cu_computer_use_status_fn(**tool_input)
                elif tool_name == "broker_run_task":
                    # Show progress message while waiting
                    connector = tool_input.get("connector", "web_generic")
                    action = tool_input.get("action", "extract")
                    progress_icons = {"cedam": "🚗", "web_generic": "🌐", "desktop_generic": "🖥️"}
                    icon = progress_icons.get(connector, "⚙️")
                    await cl.Message(
                        content=f"{icon} Running task on local agent (`{connector}` / `{action}`)... ⏳",
                        author="Alex 🤖",
                    ).send()
                    result = await _cu_run_task_async(
                        connector=connector,
                        action=action,
                        params=tool_input.get("params", {}),
                        agent_id=tool_input.get("agent_id"),
                        timeout=tool_input.get("timeout", 120),
                        credentials=tool_input.get("credentials", {}),
                    )

                # ── Output file management ────────────────────────────────
                elif tool_name == "broker_list_output_files":
                    _sort_by = tool_input.get("sort_by", "date")
                    _filter_ext = tool_input.get("filter_ext", "all").lower()
                    _exts = {".pdf", ".txt", ".xlsx", ".docx"} if _filter_ext == "all" \
                        else {f".{_filter_ext.lstrip('.')}"}
                    _files = [
                        f for f in OUTPUT_DIR.iterdir()
                        if f.is_file() and f.suffix.lower() in _exts
                    ]
                    if _sort_by == "name":
                        _files.sort(key=lambda f: f.name.lower())
                    else:
                        _files.sort(key=lambda f: f.stat().st_mtime, reverse=True)

                    if not _files:
                        result = "📂 No generated files in the output directory."
                    else:
                        _total_mb = sum(f.stat().st_size for f in _files) / (1024 * 1024)
                        _lines = [
                            f"## 📂 Generated files ({len(_files)} files, {_total_mb:.1f} MB total)\n",
                            "| # | File | Type | Size | Date |",
                            "|---|---|---|---|---|",
                        ]
                        from datetime import datetime as _dt
                        for _i, _f in enumerate(_files[:50], 1):
                            _sz = _f.stat().st_size
                            _sz_str = f"{_sz/1024:.1f} KB" if _sz < 1_000_000 else f"{_sz/1_000_000:.1f} MB"
                            _mtime = _dt.fromtimestamp(_f.stat().st_mtime).strftime("%d %b %Y %H:%M")
                            _ext = _f.suffix.upper().lstrip(".")
                            _lines.append(f"| {_i} | `{_f.name}` | {_ext} | {_sz_str} | {_mtime} |")
                        if len(_files) > 50:
                            _lines.append(f"\n_... and {len(_files) - 50} older files (showing first 50)_")
                        _lines.append(f"\n_Say **'clean up output files'** to start the selective deletion process._")
                        result = "\n".join(_lines)

                        # Trigger interactive cleanup button
                        await cl.Message(
                            content=result,
                            actions=[
                                cl.Action(
                                    name="output_cleanup_start",
                                    label="🧹 Clean up old files",
                                    payload={"total": len(_files)},
                                )
                            ],
                            author="Alex 🤖",
                        ).send()
                        # Result already sent — give empty result to agentic loop
                        result = f"[Listed {len(_files)} files. Cleanup button sent to broker.]"

                # ── Questionnaire / Form tools ─────────────────────────
                elif tool_name == "broker_send_claim_questionnaire":
                    try:
                        from main import send_claim_questionnaire_direct
                        _claim_id = tool_input.get("claim_id", "")
                        _tpl_override = tool_input.get("template_id", "")
                        _r = await send_claim_questionnaire_direct(_claim_id, _tpl_override)
                        result = json.dumps(_r, ensure_ascii=False, default=str)
                    except Exception as _eq:
                        result = f"Error sending questionnaire: {_eq}"

                elif tool_name == "broker_auto_send_questionnaires":
                    try:
                        from main import api_auto_send_questionnaires
                        _r = await api_auto_send_questionnaires()
                        result = json.dumps(_r, ensure_ascii=False, default=str)
                    except Exception as _eq:
                        result = f"Error auto-sending questionnaires: {_eq}"

                elif tool_name == "broker_check_form_status":
                    try:
                        from main import api_forms_status_summary
                        _r = await api_forms_status_summary()
                        result = json.dumps(_r, ensure_ascii=False, default=str)
                    except Exception as _eq:
                        result = f"Error checking form status: {_eq}"

                elif tool_name == "broker_run_form_reminders":
                    try:
                        from main import api_run_form_followup
                        _r = await api_run_form_followup()
                        result = json.dumps(_r, ensure_ascii=False, default=str)
                    except Exception as _eq:
                        result = f"Error running form reminders: {_eq}"

                elif tool_name == "broker_form_daily_report":
                    try:
                        from main import api_form_daily_report
                        _r = await api_form_daily_report()
                        result = json.dumps(_r, ensure_ascii=False, default=str)
                    except Exception as _eq:
                        result = f"Error generating daily report: {_eq}"

                elif tool_name == "broker_oracle_dashboard":
                    try:
                        from main import oracle_status, oracle_dashboard
                        _status = await oracle_status()
                        _dash = await oracle_dashboard()
                        result = json.dumps({"connection": _status, "dashboard": _dash}, ensure_ascii=False, default=str)
                    except Exception as _eq:
                        result = f"Error connecting to Oracle: {_eq}"

                elif tool_name == "broker_oracle_query":
                    try:
                        from main import oracle_query
                        _sql = tool_input.get("sql", "")
                        _r = await oracle_query({"sql": _sql})
                        result = json.dumps(_r, ensure_ascii=False, default=str)
                    except Exception as _eq:
                        result = f"Error executing Oracle query: {_eq}"

                elif tool_name == "broker_execute_command":
                    try:
                        from main import execute_command
                        _cmd_r = await execute_command({"command": tool_input.get("command", ""), "ref": tool_input.get("ref", "")})
                        result = json.dumps(_cmd_r, ensure_ascii=False, default=str)
                    except Exception as _eq:
                        result = f"Error executing command: {_eq}"

                elif tool_name == "broker_list_forms":
                    try:
                        import httpx
                        async with httpx.AsyncClient() as _hc:
                            _fr = await _hc.get("http://localhost:8080/api/forms/submissions", timeout=10)
                            result = _fr.text
                    except Exception as _eq:
                        result = f"Error listing forms: {_eq}"

                elif tool_name == "broker_send_form_link":
                    try:
                        import httpx
                        async with httpx.AsyncClient() as _hc:
                            _fr = await _hc.post("http://localhost:8080/api/forms/send-link", json=tool_input, timeout=15)
                            result = _fr.text
                    except Exception as _eq:
                        result = f"Error sending form link: {_eq}"

                # ── Email send: ask broker to confirm/change recipient ────
                elif tool_name == "broker_send_offer_email":
                    # Look up client email from DB to show in prompt
                    _offer_id_e = tool_input.get("offer_id", "")
                    _to_email_e = tool_input.get("to_email", "")
                    if not _to_email_e:
                        try:
                            from shared.db import get_conn as _gcn_e
                            _dc_e = _gcn_e()
                            _row_e = _dc_e.execute(
                                """SELECT c.email FROM offers o
                                   JOIN clients c ON o.client_id = c.id
                                   WHERE o.id = ?""",
                                (_offer_id_e,)
                            ).fetchone()
                            _dc_e.close()
                            _to_email_e = _row_e["email"] if _row_e else ""
                        except Exception:
                            _to_email_e = ""

                    # Ask broker to confirm or change the email address
                    _email_ask = await cl.AskUserMessage(
                        content=(
                            f"📧 **Which email address should the offer be sent to?**\n\n"
                            f"Address from client file: **{_to_email_e or 'unknown'}**\n\n"
                            f"Confirm with **yes** / **ok** to use the address above, "
                            f"or type a different email address."
                        ),
                        timeout=120,
                        author="Alex 🤖",
                    ).send()

                    if _email_ask:
                        _answer = _email_ask["output"].strip()
                        # If broker typed a valid email, use it; otherwise keep default
                        if "@" in _answer and "." in _answer.split("@")[-1]:
                            tool_input["to_email"] = _answer
                        elif _to_email_e:
                            tool_input["to_email"] = _to_email_e
                        result = execute_tool(tool_name, tool_input)
                    else:
                        result = "⚠️ Email sending was cancelled (timeout)."

                else:
                    result = execute_tool(tool_name, tool_input)

                step.output = result[:1000] + ("…" if len(result) > 1000 else "")

            # ── Audit log ──────────────────────────────────────────────────
            if ADMIN_ENABLED and user_meta.get("user_id"):
                log_audit(
                    user_id=user_meta["user_id"],
                    company_id=user_meta.get("company_id"),
                    tool_name=tool_name,
                    input_summary=str(tool_input)[:200],
                    success=True,
                    tokens=0,
                )

            # Truncate large tool results in history to avoid oversized context
            result_for_history = result[:600] + "\n[...truncated for context...]" if len(result) > 600 else result

            tool_results.append({
                "type": "tool_result",
                "tool_use_id": tool_use_id,
                "content": result_for_history,
            })

            # Attach offer file + export options when offer is created
            if tool_name == "broker_create_offer":
                output_files = sorted(OUTPUT_DIR.glob("*.txt"), key=lambda f: f.stat().st_mtime, reverse=True)
                if output_files:
                    latest = output_files[0]
                    offer_content = latest.read_text(encoding="utf-8")
                    base_title = latest.stem

                    cl.user_session.set("last_offer_file", str(latest))
                    cl.user_session.set("last_offer_content", offer_content)
                    cl.user_session.set("last_offer_title", base_title)
                    cl.user_session.set("offer_approved", False)

                    # Store tool_results in session so action callbacks can flush them
                    cl.user_session.set("_offer_pending_tool_results", list(tool_results))

                    offer_md = offer_content  # already clean markdown

                    # Display offer as nicely formatted markdown
                    await cl.Message(
                        content=offer_md,
                        author="Alex \U0001f916"
                    ).send()

                    # Persistent action buttons (no timeout — works reliably on Cloud Run)
                    await cl.Message(
                        content="**What should we do with the offer?** Choose an action:",
                        actions=[
                            cl.Action(name="offer_approve", label="\u2705 Send via Email", payload={"value": "approve"}),
                            cl.Action(name="offer_save_dashboard", label="\U0001f4ec Save to Dashboard", payload={"value": "save_dashboard"}),
                            cl.Action(name="offer_download", label="\U0001f4e5 Download (PDF / XLSX / DOCX)", payload={"value": "download"}),
                            cl.Action(name="offer_edit", label="\u270f\ufe0f Edit Offer", payload={"value": "edit"}),
                        ],
                        author="Alex \U0001f916"
                    ).send()

                    # Flush tool_results and return — action callbacks handle the rest
                    if tool_results:
                        history.append({"role": "user", "content": tool_results})
                    history.append({"role": "assistant", "content": "Offer generated. Waiting for broker action."})
                    cl.user_session.set("history", history)
                    return  # action_callback handlers take over

            # Attach export for reports
            elif tool_name in ("broker_asf_summary", "broker_bafin_summary"):
                report_type = "ASF" if tool_name == "broker_asf_summary" else "BaFin"
                base_title = f"{report_type}_Report_{date.today().isoformat()}"
                await send_export_files(result, base_title)

        # Append tool results as next user message (Anthropic format)
        if tool_results:
            history.append({"role": "user", "content": tool_results})

    # ── Send final response ────────────────────────────────────────────────
    cl.user_session.set("history", history)

    # Persist conversation history to DB (every conversation with an ID)
    if conv_id and ADMIN_ENABLED:
        try:
            save_conversation_history(conv_id, history)
        except Exception:
            pass  # non-fatal

    if final_text and final_text.strip():
        await cl.Message(content=final_text.strip(), author="Alex 🤖").send()
    # else: tool already sent output (offer file, export, etc.) — no generic message needed

    # ── Auto-improve conversation title after first full exchange ────────────
    # After the first assistant response, generate a smarter title from both sides
    _title_improved = cl.user_session.get("title_improved", False)
    if conv_id and ADMIN_ENABLED and not _title_improved and len(history) >= 2:
        try:
            # Build a 1-shot title from first user msg + first assistant reply
            _first_user = next((m["content"] for m in history if m["role"] == "user"), "")
            _first_asst = next((
                (c["text"] if isinstance(c, dict) else c)
                for m in history if m["role"] == "assistant"
                for c in (m["content"] if isinstance(m["content"], list) else [m["content"]])
                if (isinstance(c, dict) and c.get("type") == "text") or isinstance(c, str)
            ), "")
            if isinstance(_first_user, list):
                _first_user = " ".join(b.get("text", "") for b in _first_user if isinstance(b, dict) and b.get("type") == "text")
            _snippet_u = str(_first_user)[:200]
            _snippet_a = str(_first_asst)[:200]
            if _snippet_u or _snippet_a:
                _title_resp = anthropic.Anthropic().messages.create(
                    model="claude-haiku-4-5",
                    max_tokens=30,
                    messages=[{
                        "role": "user",
                        "content": (
                            f"Generează un titlu scurt (max 6 cuvinte, fără ghilimele) pentru această conversație:\n"
                            f"User: {_snippet_u}\nAlex: {_snippet_a}\n"
                            f"Titlul trebuie să fie în română, concis, descriptiv (ex: 'Ofertă CASCO Ionescu', 'Reînnoire RCA Maria Popescu', 'Raport ASF martie')."
                        )
                    }]
                )
                _smart_title = _title_resp.content[0].text.strip().strip('"\'').strip()
                if _smart_title and len(_smart_title) <= 80:
                    update_conversation_title(conv_id, _smart_title)
                    cl.user_session.set(_SK_TITLE_SET, True)
        except Exception:
            pass  # non-fatal — keep the fallback title
        cl.user_session.set("title_improved", True)
    # ────────────────────────────────────────────────────────────────────────

    # ── "Save conversation" nudge (once per session, after 3+ exchanges, no client linked) ──
    # len(history) >= 6: 3 user + 3 assistant messages = 3 full exchanges
    nudge_shown = cl.user_session.get(_SK_SAVE_NUDGE, False)
    if (
        ADMIN_ENABLED
        and conv_id
        and not client_id               # not yet linked to a client
        and not nudge_shown             # show only once per session
        and len(history) >= 6           # at least 3 full exchanges
    ):
        cl.user_session.set(_SK_SAVE_NUDGE, True)
        await cl.Message(
            content="",
            actions=[
                cl.Action(
                    name="save_conv_pick_client",
                    label="💾 Save this conversation to a client",
                    payload={"conv_id": conv_id},
                )
            ],
            author="Alex 🤖",
        ).send()
